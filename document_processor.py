import os
import re
import json
import threading
from datetime import datetime
from docx import Document
import tkinter as tk
from tkinter import filedialog, ttk

# 直接导入所有需要的模块，确保PyInstaller能够正确打包
import docx
from openpyxl import Workbook, load_workbook
from docx2pdf import convert
from PyPDF2 import PdfMerger

# 设置标志位为True，因为我们现在直接导入了这些模块
EXCEL_PROCESSING_AVAILABLE = True
PDF_CONVERSION_AVAILABLE = True
PDF_MERGING_AVAILABLE = True

class DocumentProcessor:
    def __init__(self):
        """
        初始化文档处理器
        """
        self.placeholders = set()  # 存储所有占位符
        self.user_inputs = {}  # 存储用户输入
        self.template_files = []  # 存储选中的模板文件
        self.progress_callback = None  # 进度回调函数

    def set_progress_callback(self, callback):
        """
        设置进度回调函数
        :param callback: 回调函数
        """
        self.progress_callback = callback

    def extract_placeholders_from_docx(self, file_path):
        """
        从Word文档中提取占位符
        :param file_path: Word文档路径
        :return: 占位符集合
        """
        placeholders = set()
        doc = Document(file_path)
        
        # 提取段落中的占位符
        for paragraph in doc.paragraphs:
            placeholders.update(self.find_placeholders_in_text(paragraph.text))
        
        # 提取表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    placeholders.update(self.find_placeholders_in_text(cell.text))
        
        return placeholders

    def extract_placeholders_from_xlsx(self, file_path):
        """
        从Excel文件中提取占位符
        :param file_path: Excel文件路径
        :return: 占位符集合
        """
        if not EXCEL_PROCESSING_AVAILABLE:
            raise Exception("Excel处理功能不可用，请安装openpyxl库")
        
        placeholders = set()
        workbook = load_workbook(file_path)
        
        # 遍历所有工作表
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            
            # 遍历所有单元格
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        placeholders.update(self.find_placeholders_in_text(cell.value))
        
        return placeholders

    def find_placeholders_in_text(self, text):
        """
        在文本中查找占位符
        :param text: 要搜索的文本
        :return: 占位符集合
        """
        # 使用正则表达式查找形如 {占位符名称} 的占位符
        pattern = r'\{([^}]+)\}'
        matches = re.findall(pattern, text)
        return set(matches)

    def collect_all_placeholders(self, template_files):
        """
        从所有模板文件中收集占位符
        :param template_files: 模板文件列表
        :return: 所有占位符的集合和占位符到文件的映射
        """
        all_placeholders = set()
        placeholder_files = {}  # 记录每个占位符出现在哪些文件中
        for file_path in template_files:
            if file_path.endswith('.docx'):
                placeholders = self.extract_placeholders_from_docx(file_path)
            elif file_path.endswith('.xlsx'):
                placeholders = self.extract_placeholders_from_xlsx(file_path)
            else:
                continue
                
            all_placeholders.update(placeholders)
            
            # 记录每个占位符出现的文件
            for placeholder in placeholders:
                if placeholder not in placeholder_files:
                    placeholder_files[placeholder] = []
                placeholder_files[placeholder].append(file_path)
                
        return all_placeholders, placeholder_files

    def replace_placeholders_in_docx(self, template_path, output_path, replacements):
        """
        在Word文档中替换占位符
        :param template_path: 模板文件路径
        :param output_path: 输出文件路径
        :param replacements: 替换字典
        """
        doc = Document(template_path)
        
        # 替换段落中的占位符
        for paragraph in doc.paragraphs:
            self.replace_text_in_paragraph(paragraph, replacements)
        
        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.replace_text_in_paragraph(paragraph, replacements)
        
        # 替换页眉页脚中的占位符
        for section in doc.sections:
            # 替换页眉中的占位符
            for paragraph in section.header.paragraphs:
                self.replace_text_in_paragraph(paragraph, replacements)
            
            # 替换页脚中的占位符
            for paragraph in section.footer.paragraphs:
                self.replace_text_in_paragraph(paragraph, replacements)
        
        # 保存新文档
        doc.save(output_path)

    def replace_placeholders_in_xlsx(self, template_path, output_path, replacements):
        """
        在Excel文件中替换占位符
        :param template_path: 模板文件路径
        :param output_path: 输出文件路径
        :param replacements: 替换字典
        """
        if not EXCEL_PROCESSING_AVAILABLE:
            raise Exception("Excel处理功能不可用，请安装openpyxl库")
        
        workbook = load_workbook(template_path)
        
        # 遍历所有工作表
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            
            # 遍历所有单元格
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value:
                        # 处理字符串类型的单元格
                        if isinstance(cell.value, str):
                            new_value = cell.value
                            for key, value in replacements.items():
                                placeholder = '{' + key + '}'
                                if placeholder in new_value:
                                    new_value = new_value.replace(placeholder, str(value))
                            cell.value = new_value
                        # 处理公式中的占位符（如果有的话）
                        elif isinstance(cell.value, (int, float)) and cell.data_type == 'f' and cell.value is not None:
                            # 处理公式类型
                            pass  # 公式中的占位符处理需要特殊处理，暂时跳过
        
        # 保存新文件
        workbook.save(output_path)

    def replace_text_in_paragraph(self, paragraph, replacements):
        """
        在段落中替换文本（改进版，更好地保持格式）
        :param paragraph: 段落对象
        :param replacements: 替换字典
        """
        # 遍历所有需要替换的占位符
        for placeholder, replacement in replacements.items():
            # 使用完整的占位符格式进行替换 {占位符}
            full_placeholder = f'{{{placeholder}}}'
            
            # 遍历段落中的每个run
            for run in paragraph.runs:
                if full_placeholder in run.text:
                    # 直接在run中替换文本，保持该run的所有格式属性
                    run.text = run.text.replace(full_placeholder, str(replacement))

    def process_templates(self, template_files, user_inputs, output_dir="docs"):
        """
        处理模板文件
        :param template_files: 模板文件列表
        :param user_inputs: 用户输入字典
        :param output_dir: 输出目录
        :return: 生成的文件路径列表
        """
        # 添加日期字段（如果用户没有自定义日期，则使用当天日期）
        if '日期' not in user_inputs:
            today = datetime.now()
            user_inputs['日期'] = today.strftime('%Y年%m月%d日')
        
        # 确保输出目录存在
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # 存储生成的文件路径
        generated_files = []
        
        # 为每个模板文件生成新文档
        for template_file in template_files:
            # 生成输出文件名
            base_name = os.path.basename(template_file)
            name, ext = os.path.splitext(base_name)
            output_file = f"{name}_已填充{ext}"
            output_path = os.path.join(output_dir, output_file)
            
            # 根据文件类型处理
            if template_file.endswith('.docx'):
                self.replace_placeholders_in_docx(template_file, output_path, user_inputs)
            elif template_file.endswith('.xlsx'):
                self.replace_placeholders_in_xlsx(template_file, output_path, user_inputs)
            
            generated_files.append(output_path)
            print(f"已生成文件: {output_path}")
        
        return generated_files

    def convert_docx_to_pdf(self, docx_paths, status_callback=None):
        """
        将Word文档转换为PDF
        :param docx_paths: Word文档路径列表
        :param status_callback: 状态更新回调函数
        :return: 生成的PDF文件路径列表
        """
        if not PDF_CONVERSION_AVAILABLE:
            raise Exception("PDF转换功能不可用，请安装docx2pdf库")
        
        pdf_files = []
        for docx_path in docx_paths:
            # 生成PDF文件名
            base_name = os.path.basename(docx_path)
            name, _ = os.path.splitext(base_name)
            pdf_file = os.path.join("docs", f"{name}.pdf")
            
            # 确保输出目录存在
            pdf_dir = os.path.dirname(pdf_file)
            if not os.path.exists(pdf_dir):
                os.makedirs(pdf_dir)
            
            # 更新进度窗口状态
            if self.progress_callback:
                self.progress_callback(base_name, "converting")
            
            # 尝试多种方法转换为PDF
            conversion_success = False
            
            # 方法1: 使用win32com.client方法（首选，支持WPS和Office）
            try:
                status_msg = f"正在转换: {base_name}"
                if status_callback:
                    status_callback(status_msg)
                print(f"正在尝试使用win32com.client转换: {docx_path}")
                
                import win32com.client
                import pythoncom
                
                # 初始化COM线程
                pythoncom.CoInitialize()
                
                # 添加超时机制和更好的资源管理
                word = None
                doc = None
                try:
                    # 优先尝试WPS
                    try:
                        word = win32com.client.Dispatch("KWPS.Application")
                        print("使用WPS进行转换")
                    except:
                        # 如果WPS不可用，尝试使用Microsoft Word
                        try:
                            word = win32com.client.Dispatch("Word.Application")
                            print("使用Microsoft Word进行转换")
                        except:
                            raise Exception("未找到可用的Word处理程序（WPS或Microsoft Word）")
                    
                    word.Visible = False  # 正确的属性名（大写V）
                    word.DisplayAlerts = False  # 禁用警告对话框
                    
                    # 打开文档（以只读模式）
                    doc = word.Documents.Open(os.path.abspath(docx_path), ReadOnly=True)
                    
                    # 保存为PDF
                    doc.SaveAs(os.path.abspath(pdf_file), FileFormat=17)  # 17表示PDF格式
                    
                    conversion_success = True
                except Exception as com_error:
                    error_msg = f"win32com.client转换过程中出错: {str(com_error)}"
                    print(error_msg)
                    raise com_error
                finally:
                    # 确保正确释放资源
                    try:
                        if doc:
                            doc.Close(SaveChanges=0)  # 0表示不保存更改直接关闭
                    except:
                        pass
                    try:
                        if word:
                            word.Quit()
                    except:
                        pass
                    
                    # 清理COM资源
                    pythoncom.CoUninitialize()
                
                if conversion_success and os.path.exists(pdf_file):
                    pdf_files.append(pdf_file)
                    
                    status_msg = f"已转换为PDF: {name}.pdf"
                    if status_callback:
                        status_callback(status_msg)
                    print(f"已使用win32com.client转换为PDF: {pdf_file}")
                    
                    # 更新进度窗口状态
                    if self.progress_callback:
                        self.progress_callback(base_name, "completed")
                else:
                    raise Exception("win32com.client未能生成PDF文件")
                
                conversion_success = True
            except Exception as e1:
                error_msg = f"使用win32com.client转换PDF时出错: {str(e1)}"
                print(error_msg)
                
                # 更新进度窗口状态
                if self.progress_callback:
                    self.progress_callback(base_name, "failed")
            
            # 方法2: 使用docx2pdf库作为备选方案
            if not conversion_success:
                try:
                    status_msg = f"正在转换: {base_name} (docx2pdf)"
                    if status_callback:
                        status_callback(status_msg)
                    print("正在尝试使用docx2pdf转换...")
                    
                    convert(docx_path, pdf_file)
                    pdf_files.append(pdf_file)
                    
                    status_msg = f"已转换为PDF: {name}.pdf"
                    if status_callback:
                        status_callback(status_msg)
                    print(f"已转换为PDF: {pdf_file}")
                    
                    # 更新进度窗口状态
                    if self.progress_callback:
                        self.progress_callback(base_name, "completed")
                    
                    conversion_success = True
                except Exception as e2:
                    error_msg = f"使用docx2pdf转换时出错: {str(e2)}"
                    print(error_msg)
                    
                    # 更新进度窗口状态
                    if self.progress_callback:
                        self.progress_callback(base_name, "failed")
            
            # 如果所有方法都失败了，抛出异常
            if not conversion_success:
                raise Exception("所有PDF转换方法都失败了，请检查系统配置")
        
        return pdf_files

    def convert_xlsx_to_pdf(self, xlsx_paths, status_callback=None):
        """
        将Excel文件转换为PDF
        :param xlsx_paths: Excel文件路径列表
        :param status_callback: 状态更新回调函数
        :return: 生成的PDF文件路径列表
        """
        if not PDF_CONVERSION_AVAILABLE:
            raise Exception("PDF转换功能不可用，请安装docx2pdf库")
        
        pdf_files = []
        for xlsx_path in xlsx_paths:
            # 生成PDF文件名
            base_name = os.path.basename(xlsx_path)
            name, _ = os.path.splitext(base_name)
            pdf_file = os.path.join("docs", f"{name}.pdf")
            
            # 确保输出目录存在
            pdf_dir = os.path.dirname(pdf_file)
            if not os.path.exists(pdf_dir):
                os.makedirs(pdf_dir)
            
            # 更新进度窗口状态
            if self.progress_callback:
                self.progress_callback(base_name, "converting")
            
            # 尝试多种方法转换为PDF
            conversion_success = False
            
            # 方法1: 使用win32com.client方法（首选，支持WPS和Office）
            try:
                status_msg = f"正在转换: {base_name}"
                if status_callback:
                    status_callback(status_msg)
                print(f"正在尝试使用win32com.client转换Excel: {xlsx_path}")
                
                import win32com.client
                import pythoncom
                
                # 初始化COM线程
                pythoncom.CoInitialize()
                
                # 添加更好的资源管理
                excel = None
                workbook = None
                try:
                    # 优先尝试WPS
                    try:
                        excel = win32com.client.Dispatch("KET.Application")
                        print("使用WPS表格进行转换")
                    except:
                        # 如果WPS不可用，尝试使用Microsoft Excel
                        try:
                            excel = win32com.client.Dispatch("Excel.Application")
                            print("使用Microsoft Excel进行转换")
                        except:
                            raise Exception("未找到可用的Excel处理程序（WPS表格或Microsoft Excel）")
                    
                    excel.Visible = False  # 正确的属性名（大写V）
                    excel.DisplayAlerts = False  # 禁用警告对话框
                    
                    # 打开工作簿（以只读模式）
                    workbook = excel.Workbooks.Open(os.path.abspath(xlsx_path), ReadOnly=True)
                    
                    # 导出为PDF
                    workbook.ExportAsFixedFormat(0, os.path.abspath(pdf_file))  # 0表示PDF格式
                    
                    conversion_success = True
                except Exception as com_error:
                    error_msg = f"win32com.client转换Excel过程中出错: {str(com_error)}"
                    print(error_msg)
                    raise com_error
                finally:
                    # 确保正确释放资源
                    try:
                        if workbook:
                            workbook.Close(SaveChanges=0)  # 0表示不保存更改直接关闭
                    except:
                        pass
                    try:
                        if excel:
                            excel.Quit()
                    except:
                        pass
                    
                    # 清理COM资源
                    pythoncom.CoUninitialize()
                
                if conversion_success and os.path.exists(pdf_file):
                    pdf_files.append(pdf_file)
                    
                    status_msg = f"已转换为PDF: {name}.pdf"
                    if status_callback:
                        status_callback(status_msg)
                    print(f"已使用win32com.client转换Excel为PDF: {pdf_file}")
                    
                    # 更新进度窗口状态
                    if self.progress_callback:
                        self.progress_callback(base_name, "completed")
                else:
                    raise Exception("win32com.client未能生成PDF文件")
                
                conversion_success = True
            except Exception as e1:
                error_msg = f"使用win32com.client转换Excel为PDF时出错: {str(e1)}"
                print(error_msg)
                
                # 更新进度窗口状态
                if self.progress_callback:
                    self.progress_callback(base_name, "failed")
            
            # 方法2: 使用docx2pdf库作为备选方案
            if not conversion_success:
                try:
                    status_msg = f"正在转换: {base_name} (docx2pdf)"
                    if status_callback:
                        status_callback(status_msg)
                    print("正在尝试使用docx2pdf转换Excel...")
                    
                    convert(xlsx_path, pdf_file)
                    pdf_files.append(pdf_file)
                    
                    status_msg = f"已转换为PDF: {name}.pdf"
                    if status_callback:
                        status_callback(status_msg)
                    print(f"已使用docx2pdf转换Excel为PDF: {pdf_file}")
                    
                    # 更新进度窗口状态
                    if self.progress_callback:
                        self.progress_callback(base_name, "completed")
                    
                    conversion_success = True
                except Exception as e2:
                    error_msg = f"使用docx2pdf转换Excel为PDF时出错: {str(e2)}"
                    print(error_msg)
                    
                    # 更新进度窗口状态
                    if self.progress_callback:
                        self.progress_callback(base_name, "failed")
            
            # 如果所有方法都失败了，抛出异常
            if not conversion_success:
                raise Exception("所有PDF转换方法都失败了，请检查系统配置")
        
        return pdf_files

    def cleanup_single_pdfs(self, pdf_paths, status_callback=None):
        """
        清理单个PDF文件
        :param pdf_paths: 要删除的PDF文件路径列表
        :param status_callback: 状态更新回调函数
        """
        for pdf_path in pdf_paths:
            try:
                if os.path.exists(pdf_path):
                    os.remove(pdf_path)
                    status_msg = f"已删除临时文件: {os.path.basename(pdf_path)}"
                    if status_callback:
                        status_callback(status_msg)
                    print(f"已删除临时文件: {pdf_path}")
            except Exception as e:
                error_msg = f"删除文件 {os.path.basename(pdf_path)} 时出错: {str(e)}"
                if status_callback:
                    status_callback(error_msg)
                print(f"删除文件 {pdf_path} 时出错: {str(e)}")

    def merge_pdfs(self, pdf_paths, output_path, status_callback=None):
        """
        合并多个PDF文件
        :param pdf_paths: PDF文件路径列表
        :param output_path: 输出文件路径
        :param status_callback: 状态更新回调函数
        """
        if not PDF_MERGING_AVAILABLE:
            raise Exception("PDF合并功能不可用，请安装PyPDF2库")
        
        try:
            merger = PdfMerger()
            
            for pdf_path in pdf_paths:
                try:
                    # 检查文件是否存在
                    if not os.path.exists(pdf_path):
                        raise FileNotFoundError(f"PDF文件不存在: {pdf_path}")
                    
                    merger.append(pdf_path)
                except Exception as e:
                    error_msg = f"添加PDF文件 {os.path.basename(pdf_path)} 时出错: {str(e)}"
                    if status_callback:
                        status_callback(error_msg)
                    print(f"添加PDF文件 {pdf_path} 时出错: {str(e)}")
                    raise e
            
            # 确保输出目录存在
            output_dir = os.path.dirname(output_path)
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
            
            # 合并PDF文件
            merger.write(output_path)
            merger.close()
            
            status_msg = f"已合并PDF: {os.path.basename(output_path)}"
            if status_callback:
                status_callback(status_msg)
            print(f"已合并PDF: {output_path}")
            
        except Exception as e:
            error_msg = f"合并PDF时出错: {str(e)}"
            if status_callback:
                status_callback(error_msg)
            print(f"合并PDF时出错: {str(e)}")
            raise e


class DocumentProcessorUI:
    def __init__(self, root):
        """
        初始化用户界面
        :param root: Tkinter根窗口
        """
        self.root = root
        # 将窗口居中显示
        self.center_window()
        self.root.title("填单助手 By:www.52pojie.cn@xianyuwangyou")
        self.root.geometry("800x650")
        self.root.resizable(False, False)
       
        self.processor = DocumentProcessor()
        self.template_files = []
        self.placeholders = set()
        self.placeholder_files = {}  # 存储占位符和文件的映射关系
        self.ordered_placeholders = []  # 存储有序的占位符列表
        self.current_scheme = None  # 当前选择的方案
        self.output_dir = self.load_last_output_dir()  # 输出目录，默认从配置加载
        
        self.setup_ui()
    
    def center_window(self):
        """
        将窗口居中显示在屏幕中央
        """
        # 使用固定的窗口尺寸
        width = 800
        height = 650
        
        # 获取屏幕尺寸
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        
        # 计算居中位置
        x = (screen_width - width) // 2
        y = (screen_height - height) // 2
        
        # 设置窗口位置和尺寸
        self.root.geometry(f"{width}x{height}+{x}+{y}")
    
    def center_dialog(self, dialog, width, height):
        """
        将对话框居中显示在主窗口中央
        :param dialog: 对话框窗口
        :param width: 对话框宽度
        :param height: 对话框高度
        """
        # 获取主窗口位置和尺寸
        main_x = self.root.winfo_x()
        main_y = self.root.winfo_y()
        main_width = self.root.winfo_width()
        main_height = self.root.winfo_height()
        
        # 计算居中位置（相对于主窗口）
        x = main_x + (main_width - width) // 2
        y = main_y + (main_height - height) // 2
        
        # 确保对话框不会显示在屏幕外
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        
        # 调整x坐标，确保对话框在屏幕内
        if x < 0:
            x = 0
        elif x + width > screen_width:
            x = screen_width - width
            
        # 调整y坐标，确保对话框在屏幕内
        if y < 0:
            y = 0
        elif y + height > screen_height:
            y = screen_height - height
        
        # 设置对话框位置和尺寸
        dialog.geometry(f"{width}x{height}+{x}+{y}")
    
    def set_dialog_icon(self, dialog):
        """
        设置对话框图标，直接读取根目录下的icon.ico文件
        :param dialog: 对话框窗口
        """
        try:
            # 直接读取根目录下的icon.ico文件
            root_dir = os.path.dirname(os.path.abspath(__file__))
            icon_path = os.path.join(root_dir, 'icon.ico')
            
            # 如果图标文件存在，则设置对话框图标
            if os.path.exists(icon_path):
                dialog.iconbitmap(icon_path)
        except Exception as e:
            # 如果设置图标失败，不中断程序运行
            print(f"设置对话框图标失败: {e}")
            pass
    
    def load_last_output_dir(self):
        """
        加载上次使用的输出目录
        :return: 上次使用的输出目录路径
        """
        try:
            if os.path.exists("app_data.json"):
                with open("app_data.json", "r", encoding="utf-8") as f:
                    data = json.load(f)
                return data.get("config", {}).get("last_output_dir", "docs")
            else:
                return "docs"
        except Exception as e:
            print(f"加载配置文件时出错: {e}")
            return "docs"
    
    def load_last_template_dir(self):
        """
        加载上次使用的模板目录
        :return: 上次使用的模板目录路径
        """
        try:
            if os.path.exists("app_data.json"):
                with open("app_data.json", "r", encoding="utf-8") as f:
                    data = json.load(f)
                return data.get("config", {}).get("last_template_dir", "docs")
            else:
                return "docs"
        except Exception as e:
            print(f"加载配置文件时出错: {e}")
            return "docs"
    
    def save_last_output_dir(self, output_dir):
        """
        保存当前使用的输出目录到配置文件
        :param output_dir: 输出目录
        """
        try:
            # 读取现有配置
            config = {}
            if os.path.exists("app_data.json"):
                with open("app_data.json", "r", encoding="utf-8") as f:
                    config = json.load(f)
            
            # 确保config键存在
            if "config" not in config:
                config["config"] = {}
            
            # 更新输出目录
            config["config"]["last_output_dir"] = output_dir
            
            # 保存配置
            with open("app_data.json", "w", encoding="utf-8") as f:
                json.dump(config, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"保存配置文件时出错: {e}")

    def save_last_template_dir(self, template_dir):
        """
        保存最后使用的模板目录到配置文件
        :param template_dir: 模板目录路径
        """
        try:
            # 读取现有配置
            config = {}
            if os.path.exists("app_data.json"):
                with open("app_data.json", "r", encoding="utf-8") as f:
                    config = json.load(f)
            
            # 确保config键存在
            if "config" not in config:
                config["config"] = {}
            
            # 更新模板目录
            config["config"]["last_template_dir"] = template_dir
            
            # 保存配置
            with open("app_data.json", "w", encoding="utf-8") as f:
                json.dump(config, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"保存配置文件时出错: {e}")

    def get_placeholder_config(self, placeholder):
        """
        获取占位符配置
        :param placeholder: 占位符名称
        :return: 配置字典
        """
        try:
            if os.path.exists("app_data.json"):
                with open("app_data.json", "r", encoding="utf-8") as f:
                    data = json.load(f)
                return data.get("placeholder_configs", {}).get(placeholder, {})
            else:
                return {}
        except Exception as e:
            print(f"加载占位符配置时出错: {e}")
            return {}
    
    def save_placeholder_config(self, placeholder, config):
        """
        保存占位符配置
        :param placeholder: 占位符名称
        :param config: 配置字典
        """
        try:
            # 读取现有配置
            main_config = {}
            if os.path.exists("app_data.json"):
                with open("app_data.json", "r", encoding="utf-8") as f:
                    main_config = json.load(f)
            
            # 确保placeholder_configs键存在
            if "placeholder_configs" not in main_config:
                main_config["placeholder_configs"] = {}
            
            # 保存当前占位符配置
            main_config["placeholder_configs"][placeholder] = config
            
            # 保存配置
            with open("app_data.json", "w", encoding="utf-8") as f:
                json.dump(main_config, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"保存占位符配置时出错: {e}")

    def load_user_inputs_for_scheme(self, scheme_name):
        """
        加载指定方案的用户输入
        :param scheme_name: 方案名称
        :return: 用户输入字典
        """
        try:
            if os.path.exists("app_data.json"):
                with open("app_data.json", "r", encoding="utf-8") as f:
                    data = json.load(f)
                    return data.get("user_inputs", {}).get(scheme_name, {})
            else:
                return {}
        except Exception as e:
            print(f"加载用户输入时出错: {e}")
            return {}

    def save_user_inputs_for_scheme(self, scheme_name, user_inputs):
        """
        保存指定方案的用户输入
        :param scheme_name: 方案名称
        :param user_inputs: 用户输入字典
        """
        try:
            # 读取现有配置
            config = {}
            if os.path.exists("app_data.json"):
                with open("app_data.json", "r", encoding="utf-8") as f:
                    config = json.load(f)
            
            # 确保user_inputs键存在
            if "user_inputs" not in config:
                config["user_inputs"] = {}
            
            # 保存当前方案的用户输入（之前会排除日期字段，现在不再排除）
            # 修复：删除无用的filtered_inputs变量，直接使用user_inputs
            config["user_inputs"][scheme_name] = user_inputs
            
            # 保存配置
            with open("app_data.json", "w", encoding="utf-8") as f:
                json.dump(config, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"保存用户输入时出错: {e}")

    def setup_ui(self):
        """
        设置用户界面
        """

        # 创建标签页控件
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # 绑定标签页切换事件
        self.notebook.bind("<<NotebookTabChanged>>", self.on_tab_changed)
        
        # 创建主操作标签页
        self.main_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.main_frame, text="数据录入")
        
        # 创建配置方案标签页
        self.config_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.config_frame, text="配置文档组合")
        
        # 创建模板制作标签页
        self.template_maker_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.template_maker_frame, text="配置文档模板")
        
        # 创建选项标签页
        self.options_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.options_frame, text="选项设置")
        
        # 创建状态栏（提前创建，确保其他组件可以使用）
        self.status_bar = ttk.Label(self.root, text="就绪", relief=tk.SUNKEN, anchor=tk.W)
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)
        
        self.setup_main_tab()
        self.setup_config_tab()
        self.setup_template_maker_tab()
        self.setup_options_tab()

    def setup_template_maker_tab(self):
        """
        设置模板制作标签页
        """
        template_maker_frame = ttk.Frame(self.template_maker_frame, padding="10")
        template_maker_frame.pack(fill=tk.BOTH, expand=True)
        
        # 配置网格权重
        template_maker_frame.columnconfigure(0, weight=1)  # 左侧占位符列表
        template_maker_frame.columnconfigure(1, weight=2)  # 右侧文档操作区
        template_maker_frame.rowconfigure(0, weight=1)
        
        # 左侧框架：占位符列表
        left_frame = ttk.LabelFrame(template_maker_frame, text="可用占位符", padding="10")
        left_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(0, 5))
        
        # 右侧框架：文档操作区
        right_frame = ttk.LabelFrame(template_maker_frame, text="文档操作", padding="10")
        right_frame.grid(row=0, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(5, 0))
        
        # 左侧内容：占位符列表
        left_frame.columnconfigure(0, weight=1)
        left_frame.rowconfigure(1, weight=1)
        
        # 占位符操作区域（现在是空的，因为我们把按钮移到了下方）
        placeholder_button_frame = ttk.Frame(left_frame)
        placeholder_button_frame.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))
        
        # 占位符列表框
        self.placeholder_listbox = tk.Listbox(left_frame, height=15)
        self.placeholder_listbox.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))
        # 添加双击事件绑定
        self.placeholder_listbox.bind("<Double-Button-1>", self.edit_placeholder)
        # 添加选择事件绑定
        self.placeholder_listbox.bind("<<ListboxSelect>>", self.on_placeholder_select)
        
        # 初始化时添加默认的日期占位符（但标记为不可用）
        self.placeholder_listbox.insert(tk.END, "日期")
        self.placeholder_listbox.itemconfig(0, {'fg': 'gray'})
        
        # 添加新占位符按钮（移动到列表上方）
        self.add_placeholder_button_middle = ttk.Button(left_frame, text="添加新占位符", command=self.add_new_placeholder, state=tk.DISABLED)
        self.add_placeholder_button_middle.grid(row=0, column=0, columnspan=2, pady=(0, 10), sticky=(tk.W, tk.E))
        
        # 占位符列表框
        self.placeholder_listbox = tk.Listbox(left_frame, height=15)
        self.placeholder_listbox.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))
        # 添加双击事件绑定
        self.placeholder_listbox.bind("<Double-Button-1>", self.edit_placeholder)
        # 添加选择事件绑定
        self.placeholder_listbox.bind("<<ListboxSelect>>", self.on_placeholder_select)
        
        # 初始化时添加默认的日期占位符（但标记为不可用）
        self.placeholder_listbox.insert(tk.END, "日期")
        self.placeholder_listbox.itemconfig(0, {'fg': 'gray'})
        
        # 按钮区域：刷新占位符和复制占位符到剪贴板按钮（垂直排列，宽度一致）
        placeholder_buttons_frame = ttk.Frame(left_frame)
        placeholder_buttons_frame.grid(row=2, column=0, columnspan=2, pady=(0, 10), sticky=(tk.W, tk.E))
        
        self.refresh_placeholder_button_middle = ttk.Button(placeholder_buttons_frame, text="刷新占位符", command=self.refresh_placeholders, state=tk.DISABLED)
        self.refresh_placeholder_button_middle.pack(fill=tk.X, pady=(0, 5))
        
        self.copy_placeholder_button = ttk.Button(placeholder_buttons_frame, text="复制占位符到剪贴板", command=self.copy_placeholder_to_clipboard, state=tk.DISABLED)
        self.copy_placeholder_button.pack(fill=tk.X, pady=(0, 5))
        
        # 删除占位符按钮框架
        delete_button_frame = ttk.Frame(left_frame)
        delete_button_frame.grid(row=3, column=0, columnspan=2, pady=(0, 10), sticky=(tk.W, tk.E))
        
        # 添加删除占位符按钮
        self.delete_placeholder_button = ttk.Button(delete_button_frame, text="删除占位符", command=self.delete_placeholder, state=tk.DISABLED)
        self.delete_placeholder_button.pack(fill=tk.X)
        
        # 右侧内容：文档操作区域
        right_frame.columnconfigure(0, weight=1)
        right_frame.rowconfigure(1, weight=1)
        
        # 文档操作按钮
        doc_button_frame = ttk.Frame(right_frame)
        doc_button_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=(0, 10))
        
        ttk.Button(doc_button_frame, text="选择文档目录", command=self.select_and_open_folder).pack(side=tk.LEFT)

        # 转换格式为docx/xlsx
        self.convert_button = ttk.Button(doc_button_frame, text="转换为docx/xlsx", command=self.convert_wps_to_docx, state=tk.DISABLED)
        self.convert_button.pack(side=tk.LEFT, padx=(10, 0))

        # 打开模板文件夹
        self.open_folder_button = ttk.Button(doc_button_frame, text="打开文档目录", command=self.open_selected_folder, state=tk.DISABLED)
        self.open_folder_button.pack(side=tk.LEFT, padx=(10, 0))

        # 刷新文档目录
        self.refresh_folder_button = ttk.Button(doc_button_frame, text="刷新文档目录", command=self.refresh_folder_info)
        self.refresh_folder_button.pack(side=tk.LEFT, padx=(10, 0))
        
        # 文档信息显示区域
        self.doc_info_text = tk.Text(right_frame, height=15, wrap=tk.WORD)
        self.doc_info_text.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))
        self.doc_info_text.insert(tk.END, "请选择一个文件夹以查看信息")
        self.doc_info_text.config(state=tk.DISABLED)
        
        # 滚动条
        doc_info_scrollbar = ttk.Scrollbar(right_frame, orient="vertical", command=self.doc_info_text.yview)
        doc_info_scrollbar.grid(row=1, column=1, sticky=(tk.N, tk.S), pady=(0, 10))
        self.doc_info_text.configure(yscrollcommand=doc_info_scrollbar.set)
        
        # 使用说明
        instruction_frame = ttk.Frame(right_frame)
        instruction_frame.grid(row=2, column=0, sticky=(tk.W, tk.E), pady=(10, 0))

    def setup_options_tab(self):
        """
        设置选项标签页
        """
        options_frame = ttk.Frame(self.options_frame, padding="20")
        options_frame.pack(fill=tk.BOTH, expand=True)
        
        # 设置网格权重
        options_frame.columnconfigure(0, weight=1)
        options_frame.rowconfigure(0, weight=1)
        
        # 设置选项区域的网格权重
        options_frame.columnconfigure(0, weight=1)
        
        # 设置区域
        settings_frame = ttk.LabelFrame(options_frame, text="设置", padding="10")
        settings_frame.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(0, 20))
        settings_frame.columnconfigure(1, weight=1)
        
        # PDF转换线程数设置
        ttk.Label(settings_frame, text="PDF转换线程数:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.pdf_threads_var = tk.StringVar(value=self.load_pdf_threads_setting())
        pdf_threads_spinbox = ttk.Spinbox(
            settings_frame, 
            from_=1, 
            to=5, 
            textvariable=self.pdf_threads_var,
            width=5,
            command=self.save_pdf_threads_setting
        )
        pdf_threads_spinbox.grid(row=0, column=1, sticky=tk.W, padx=(10, 0), pady=5)
        
        # 保存按钮
        save_button = ttk.Button(settings_frame, text="保存设置", command=self.save_pdf_threads_setting)
        save_button.grid(row=1, column=0, sticky=tk.W, padx=(0, 10), pady=5)
                
        self.pdf_threads_var.trace('w', lambda *args: self.save_pdf_threads_setting())
        
        # 检查更新区域
        update_frame = ttk.LabelFrame(options_frame, text="软件更新", padding="10")
        update_frame.grid(row=2, column=0, sticky=(tk.W, tk.E), pady=(0, 20))
        update_frame.columnconfigure(1, weight=1)
        
        # 检查更新按钮
        self.update_button = ttk.Button(update_frame, text="检查更新")
        self.update_button.grid(row=0, column=0, padx=(0, 10), pady=5, sticky=tk.W)
        
        # 添加说明文字
        ttk.Label(update_frame, text="检查是否有新版本可用").grid(row=0, column=1, pady=5, sticky=tk.W)
        
        # 添加版本信息
        version_frame = ttk.Frame(update_frame)
        version_frame.grid(row=1, column=0, columnspan=2, pady=(10, 0), sticky=(tk.W, tk.E))
        ttk.Label(version_frame, text="当前版本: v1.1.0").pack(anchor=tk.W)
        
        # 关于区域
        about_frame = ttk.LabelFrame(options_frame, text="关于", padding="10")
        about_frame.grid(row=3, column=0, sticky=(tk.W, tk.E), pady=(0, 20))
        about_frame.columnconfigure(0, weight=1)
        
        # 关于信息
        about_text = ttk.Label(about_frame, 
                              text="填单助手是一个自动化的填单工具，"
                                   "通过替换Word和Excel模板中的占位符，"
                                   "批量生成文档。\n\n"
                                   "作者: 咸鱼网友\nBy:XianYuWangYou")
        about_text.grid(row=0, column=0, pady=(0, 10), sticky=tk.W)
        
        # 论坛链接
        link_frame = ttk.Label(about_frame)
        link_frame.grid(row=1, column=0, sticky=tk.W, padx=(0, 0))

        link_title = ttk.Label(link_frame, text="吾爱破解论坛:")
        link_title.pack(side=tk.LEFT)
        
        forum_link = ttk.Label(link_frame, text="https://www.52pojie.cn/thread-2053988-1-1.html", foreground="blue", cursor="hand2")
        forum_link.pack(side=tk.LEFT, padx=(5, 0))
        forum_link.bind("<Button-1>", lambda e: self.open_forum_link())

        # Gitee链接
        link_frame = ttk.Label(about_frame)
        link_frame.grid(row=2, column=0, sticky=tk.W, padx=(0, 0))

        link_title = ttk.Label(link_frame, text="Gitee项目:")
        link_title.pack(side=tk.LEFT)
        
        forum_link = ttk.Label(link_frame, text="https://gitee.com/xianyuwangyou/TianDanAssistant", foreground="blue", cursor="hand2")
        forum_link.pack(side=tk.LEFT, padx=(5, 0))
        forum_link.bind("<Button-1>", lambda e: self.open_forum_link())

        # Github链接
        link_frame = ttk.Label(about_frame)
        link_frame.grid(row=3, column=0, sticky=tk.W, padx=(0, 0))

        link_title = ttk.Label(link_frame, text="Github项目:")
        link_title.pack(side=tk.LEFT)
        
        forum_link = ttk.Label(link_frame, text="https://github.com/XianYuWangYou/TianDanAssistant", foreground="blue", cursor="hand2")
        forum_link.pack(side=tk.LEFT, padx=(5, 0))
        forum_link.bind("<Button-1>", lambda e: self.open_forum_link())

        # 添加空白区域以填充空间
        options_frame.rowconfigure(4, weight=1)
        spacer_frame = ttk.Frame(options_frame)
        spacer_frame.grid(row=4, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
    
    def open_forum_link(self):
        """
        打开吾爱破解论坛链接
        """
        import webbrowser
        webbrowser.open("https://www.52pojie.cn/")
    
    def load_pdf_threads_setting(self):
        """
        加载PDF转换线程数设置
        """
        try:
            if os.path.exists("app_data.json"):
                with open("app_data.json", "r", encoding="utf-8") as f:
                    data = json.load(f)
                return str(data.get("settings", {}).get("pdf_threads", 1))
            else:
                return "1"
        except Exception as e:
            print(f"加载PDF线程设置时出错: {e}")
            return "1"
    
    def save_pdf_threads_setting(self, *args):
        """
        保存PDF转换线程数设置
        """
        try:
            # 获取当前设置值
            try:
                threads_value = int(self.pdf_threads_var.get())
                # 确保值在有效范围内
                threads_value = max(1, min(5, threads_value))
                # 更新变量值（以防输入了范围外的值）
                self.pdf_threads_var.set(str(threads_value))
            except ValueError:
                # 如果转换失败，使用默认值
                threads_value = 1
                self.pdf_threads_var.set("1")
            
            # 读取现有配置
            config = {}
            if os.path.exists("app_data.json"):
                with open("app_data.json", "r", encoding="utf-8") as f:
                    config = json.load(f)
            
            # 确保settings键存在
            if "settings" not in config:
                config["settings"] = {}
            
            # 更新PDF线程数设置
            config["settings"]["pdf_threads"] = threads_value
            
            # 保存配置
            with open("app_data.json", "w", encoding="utf-8") as f:
                json.dump(config, f, ensure_ascii=False, indent=2)
                
        except Exception as e:
            print(f"保存PDF线程设置时出错: {e}")

    def ask_to_open_output_dir(self):
        """
        询问用户是否打开输出文件夹
        """
        try:
            from tkinter import messagebox
            result = messagebox.askyesno("操作完成", "文档已生成完成，是否打开保存目录查看文件？")
            if result:
                self.open_output_dir()
        except Exception as e:
            print(f"询问打开输出文件夹时出错: {e}")

    def open_output_dir(self):
        """
        打开输出文件夹
        """
        try:
            import os
            # 确保路径是绝对路径
            abs_output_dir = os.path.abspath(self.output_dir)
            # 使用系统默认方式打开文件夹
            os.startfile(abs_output_dir)
        except Exception as e:
            self.log_and_status(f"打开输出文件夹时出错: {str(e)}")

    def on_tab_changed(self, event):
        """
        标签页切换事件处理
        """
        # 获取当前选中的标签页索引
        current_tab = self.notebook.index(self.notebook.select())
        
        # 如果切换到模板制作标签页（索引为2），则刷新占位符列表
        if current_tab == 2:
            self.refresh_placeholders_list()
        # 如果切换到配置方案标签页（索引为1），刷新下拉菜单和占位符列表
        elif current_tab == 1:
            self.load_saved_schemes_combobox()
            self.refresh_placeholders_list()

    def update_status(self, message):
        """
        更新状态栏显示内容
        :param message: 要显示的状态信息
        """
        self.status_bar.config(text=message)
        self.root.update_idletasks()
    
    def log_and_status(self, message):
        """
        在控制台输出日志信息，并在状态栏显示
        :param message: 要输出和显示的信息
        """
        print(message)
        self.update_status(message)

    def select_and_open_folder(self):
        """
        选择并打开文件夹
        """
        # 获取上次使用的模板目录
        last_template_dir = self.load_last_template_dir()
        
        folder_path = filedialog.askdirectory(
            title="选择模板文件夹",
            initialdir=last_template_dir
        )
        
        if folder_path:
            # 保存所选文件夹的目录
            self.save_last_template_dir(folder_path)
            
            # 保存选择的路径供转换功能使用
            self.selected_template_folder = folder_path
            
            # 启用转换按钮
            self.convert_button.config(state=tk.NORMAL)
            
            # 启用打开目录按钮
            self.open_folder_button.config(state=tk.NORMAL)
            
            # 启用占位符相关按钮
            self.add_placeholder_button_middle.config(state=tk.NORMAL)
            self.refresh_placeholder_button_middle.config(state=tk.NORMAL)
            self.copy_placeholder_button.config(state=tk.NORMAL)
            self.delete_placeholder_button.config(state=tk.NORMAL)
            
            # 显示文件夹中的文件信息
            self.display_folder_info(folder_path)
            
            # 更新占位符列表
            self.refresh_placeholders_list()

    def open_selected_folder(self):
        """
        打开用户选择的目录文件夹
        """
        if hasattr(self, 'selected_template_folder') and self.selected_template_folder:
            try:
                os.startfile(self.selected_template_folder)
            except Exception as e:
                self.log_and_status(f"打开文件夹时出错: {str(e)}")
        else:
            self.log_and_status("请先选择文档目录")

    def refresh_folder_info(self):
        """
        刷新文档目录信息显示
        """
        # 检查是否已选择文档目录
        if hasattr(self, 'selected_template_folder') and self.selected_template_folder:
            self.display_folder_info(self.selected_template_folder)
            self.log_and_status("文档目录信息已刷新")
        else:
            self.log_and_status("请先选择文档目录")

    def display_folder_info(self, folder_path):
        """
        显示文件夹中的文件信息
        :param folder_path: 文件夹路径
        """
        try:
            # 获取文件夹中的所有文件
            all_files = []
            for file in os.listdir(folder_path):
                file_path = os.path.join(folder_path, file)
                if os.path.isfile(file_path):
                    all_files.append(file_path)
            
            # 统计各类文件数量
            doc_files = []  # .doc, .wps, .wpt
            xls_files = []  # .xls, .et
            docx_files = []  # .docx
            xlsx_files = []  # .xlsx
            other_files = []  # 其他文件
            
            for file_path in all_files:
                ext = os.path.splitext(file_path)[1].lower()
                if ext in ['.doc', '.wps', '.wpt']:
                    doc_files.append(file_path)
                elif ext in ['.xls', '.et']:
                    xls_files.append(file_path)
                elif ext == '.docx':
                    docx_files.append(file_path)
                elif ext == '.xlsx':
                    xlsx_files.append(file_path)
                else:
                    other_files.append(file_path)
            
            # 需要转换的文件数
            files_to_convert = len(doc_files) + len(xls_files)
            
            # 无法转换的文件数
            files_cannot_convert = len(other_files)
            
            # 构建显示信息
            info_lines = []
            info_lines.append(f"文件夹路径: {folder_path}")
            info_lines.append(f"文件总个数: {len(all_files)}")
            info_lines.append(f"DOC/WPS文件个数: {len(doc_files)}")
            info_lines.append(f"XLS/ET文件个数: {len(xls_files)}")
            info_lines.append(f"DOCX文件个数: {len(docx_files)}")
            info_lines.append(f"XLSX文件个数: {len(xlsx_files)}")
            info_lines.append(f"需转换文件个数: {files_to_convert}")
            info_lines.append(f"无法转换文件个数: {files_cannot_convert}")
            info_lines.append("")
            info_lines.append("文件列表:")
            
            # 添加文件列表
            for file_path in all_files:
                filename = os.path.basename(file_path)
                ext = os.path.splitext(file_path)[1].lower()
                file_type = "未知"
                if ext in ['.doc', '.wps', '.wpt']:
                    file_type = "需转换(DOC/WPS)"
                elif ext in ['.xls', '.et']:
                    file_type = "需转换(XLS/ET)"
                elif ext in ['.docx', '.xlsx']:
                    file_type = "已转换格式"
                else:
                    file_type = "无法转换"
                
                info_lines.append(f"  {filename} [{file_type}]")
            
            # 更新文档信息显示
            self.doc_info_text.config(state=tk.NORMAL)
            self.doc_info_text.delete(1.0, tk.END)
            self.doc_info_text.insert(1.0, '\n'.join(info_lines))
            self.doc_info_text.config(state=tk.DISABLED)
            
        except Exception as e:
            error_info = f"无法读取文件夹信息: {str(e)}"
            self.doc_info_text.config(state=tk.NORMAL)
            self.doc_info_text.delete(1.0, tk.END)
            self.doc_info_text.insert(1.0, error_info)
            self.doc_info_text.config(state=tk.DISABLED)

    def convert_wps_to_docx(self):
        """
        将指定文件夹内的WPS文件转换为DOCX格式（在新线程中执行）
        """
        # 在新线程中执行转换操作
        thread = threading.Thread(target=self._convert_wps_to_docx_thread)
        thread.daemon = True  # 设置为守护线程，确保主程序退出时线程也会退出
        thread.start()
    
    def _convert_wps_to_docx_thread(self):
        """
        在线程中执行WPS/ET/XLS到DOCX/XLSX的转换操作
        """
        # 检查是否已选择模板文件夹
        if not hasattr(self, 'selected_template_folder') or not self.selected_template_folder:
            self.log_and_status("请先点击\"打开模板目录\"按钮选择文件夹")
            # 禁用转换按钮
            self.root.after(0, lambda: self.convert_button.config(state=tk.DISABLED))
            return
        
        # 使用已选择的文件夹路径
        folder_path = self.selected_template_folder
        
        # 保存所选文件夹的目录
        self.save_last_template_dir(folder_path)
        
        # 查找文件夹中的所有文件
        all_files = []
        for file in os.listdir(folder_path):
            file_path = os.path.join(folder_path, file)
            if os.path.isfile(file_path):
                all_files.append(file_path)
        
        # 筛选支持转换的文件
        supported_extensions = ('.wps', '.wpt', '.doc', '.et', '.xls')
        files_to_convert = []
        files_to_move = []
        
        # 只有后缀不为docx、doc、wps、xlsx、xls、et才移动到"[未转换]"文件夹
        unconvertible_extensions = ('.docx', '.doc', '.wps', '.wpt', '.xlsx', '.xls', '.et')
        
        for file_path in all_files:
            if os.path.splitext(file_path)[1].lower() in supported_extensions:
                files_to_convert.append(file_path)
            elif os.path.splitext(file_path)[1].lower() not in unconvertible_extensions:
                files_to_move.append(file_path)
        
        # 创建"未转换"文件夹并移动不支持的文件
        if files_to_move:
            unconverted_folder = os.path.join(folder_path, "[未转换]")
            if not os.path.exists(unconverted_folder):
                os.makedirs(unconverted_folder)
            
            for file_path in files_to_move:
                try:
                    filename = os.path.basename(file_path)
                    destination = os.path.join(unconverted_folder, filename)
                    # 如果目标文件已存在，添加序号
                    counter = 1
                    base_name, ext = os.path.splitext(filename)
                    while os.path.exists(destination):
                        new_filename = f"{base_name}_{counter}{ext}"
                        destination = os.path.join(unconverted_folder, new_filename)
                        counter += 1
                    
                    os.rename(file_path, destination)
                    self.update_status(f"已移动不支持的文件: {filename}")
                except Exception as e:
                    self.log_and_status(f"移动文件 {os.path.basename(file_path)} 失败: {str(e)}")
        
        if not files_to_convert:
            self.log_and_status("在选定的文件夹中未找到支持转换的文件")
            return
        
        self.update_status("开始转换WPS/ET/XLS/DOC文件...")
        
        # 尝试使用win32com.client进行转换
        wps_app = None
        try:
            import win32com.client
            import pythoncom
            import threading
            import time
            
            # 初始化COM线程
            pythoncom.CoInitialize()
            
            success_count = 0
            # 分别处理文档文件和表格文件
            doc_files = [f for f in files_to_convert if os.path.splitext(f)[1].lower() in ['.wps', '.wpt', '.doc']]
            et_xls_files = [f for f in files_to_convert if os.path.splitext(f)[1].lower() in ['.et', '.xls']]
            
            # 处理文档文件 (WPS/DOC -> DOCX)
            if doc_files:
                # 创建WPS文字应用程序对象
                try:
                    wps_app = win32com.client.Dispatch("KWPS.Application")
                except:
                    # 如果KWPS不可用，尝试使用WPS
                    try:
                        wps_app = win32com.client.Dispatch("Ket.Application")
                    except:
                        # 如果WPS也不可用，尝试使用Microsoft Word
                        try:
                            wps_app = win32com.client.Dispatch("Word.Application")
                        except:
                            self.log_and_status("未找到可用的WPS或Word应用程序来处理文档文件")
                            wps_app = None
                
                if wps_app:
                    wps_app.Visible = False
                    wps_app.DisplayAlerts = False
                    
                    for file_path in doc_files:
                        try:
                            # 获取文件扩展名
                            file_ext = os.path.splitext(file_path)[1].lower()
                            
                            # 打开文件
                            doc = wps_app.Documents.Open(file_path)
                            
                            # 生成DOCX文件名
                            docx_file = os.path.splitext(file_path)[0] + ".docx"
                            
                            # 另存为DOCX格式 (12是DOCX格式的代码)
                            doc.SaveAs(docx_file, 12)
                            
                            # 关闭文档
                            doc.Close()
                            
                            success_count += 1
                            file_type = "WPS" if file_ext in ['.wps', '.wpt'] else "DOC"
                            self.update_status(f"已转换 {file_type} 文件: {os.path.basename(file_path)} -> {os.path.basename(docx_file)}")
                        except Exception as e:
                            file_type = "WPS" if os.path.splitext(file_path)[1].lower() in ['.wps', '.wpt'] else "DOC"
                            self.log_and_status(f"转换失败 {file_type} 文件 {os.path.basename(file_path)}: {str(e)}")
                    
                    # 关闭WPS应用程序
                    try:
                        wps_app.Quit()
                    except:
                        pass
                    wps_app = None
            
            # 处理电子表格文件 (ET/XLS -> XLSX)
            if et_xls_files:
                # 创建WPS表格应用程序对象
                try:
                    et_app = win32com.client.Dispatch("KET.Application")
                except:
                    # 如果KET不可用，尝试使用Microsoft Excel
                    try:
                        et_app = win32com.client.Dispatch("Excel.Application")
                    except:
                        self.log_and_status("未找到可用的WPS表格或Excel应用程序来处理电子表格文件")
                        et_app = None
                
                if et_app:
                    et_app.Visible = False
                    et_app.DisplayAlerts = False
                    
                    for file_path in et_xls_files:
                        try:
                            # 获取文件扩展名
                            file_ext = os.path.splitext(file_path)[1].lower()
                            
                            # 打开文件（添加超时和错误处理）
                            try:
                                workbook = et_app.Workbooks.Open(file_path)
                                
                                # 生成XLSX文件名
                                xlsx_file = os.path.splitext(file_path)[0] + ".xlsx"
                                
                                # 另存为XLSX格式 (51是XLSX格式的代码)
                                workbook.SaveAs(xlsx_file, 51)
                                
                                # 关闭工作簿
                                workbook.Close()
                                
                                success_count += 1
                                file_type = "ET" if file_ext == '.et' else "XLS"
                                self.update_status(f"已转换 {file_type} 文件: {os.path.basename(file_path)} -> {os.path.basename(xlsx_file)}")
                            except Exception as open_error:
                                file_type = "ET" if os.path.splitext(file_path)[1].lower() == '.et' else "XLS"
                                self.log_and_status(f"打开或转换失败 {file_type} 文件 {os.path.basename(file_path)}: {str(open_error)}")
                        except Exception as e:
                            file_type = "ET" if os.path.splitext(file_path)[1].lower() == '.et' else "XLS"
                            self.log_and_status(f"转换失败 {file_type} 文件 {os.path.basename(file_path)}: {str(e)}")
                    
                    # 关闭ET应用程序
                    try:
                        et_app.Quit()
                    except:
                        pass
            
            self.log_and_status(f"转换完成: 成功 {success_count}/{len(files_to_convert)} 个文件")
            
            # 询问用户是否删除源文件
            if success_count > 0:
                try:
                    from tkinter import messagebox
                    result = messagebox.askyesno("转换完成", "是否删除已成功转换的源文件？\n(选择\"否\"将把源文件移动到\"源文件\"文件夹中)")
                    if result:
                        deleted_count = 0
                        for file_path in files_to_convert:
                            try:
                                os.remove(file_path)
                                deleted_count += 1
                                self.update_status(f"已删除源文件: {os.path.basename(file_path)}")
                            except Exception as e:
                                self.log_and_status(f"删除源文件 {os.path.basename(file_path)} 失败: {str(e)}")
                        self.log_and_status(f"已删除 {deleted_count}/{len(files_to_convert)} 个源文件")
                    else:
                        # 用户选择"否"，将源文件移动到"源文件"文件夹
                        source_folder = os.path.join(os.path.dirname(files_to_convert[0]), "源文件")
                        if not os.path.exists(source_folder):
                            os.makedirs(source_folder)
                        
                        moved_count = 0
                        for file_path in files_to_convert:
                            try:
                                filename = os.path.basename(file_path)
                                destination = os.path.join(source_folder, filename)
                                # 如果目标文件已存在，添加序号
                                counter = 1
                                base_name, ext = os.path.splitext(filename)
                                while os.path.exists(destination):
                                    new_filename = f"{base_name}_{counter}{ext}"
                                    destination = os.path.join(source_folder, new_filename)
                                    counter += 1
                                
                                os.rename(file_path, destination)
                                moved_count += 1
                                self.update_status(f"已移动源文件到\"源文件\"文件夹: {filename}")
                            except Exception as e:
                                self.log_and_status(f"移动源文件 {os.path.basename(file_path)} 失败: {str(e)}")
                        self.log_and_status(f"已移动 {moved_count}/{len(files_to_convert)} 个源文件到\"源文件\"文件夹")
                except Exception as e:
                    self.log_and_status(f"处理删除/移动操作时出错: {str(e)}")
            
        except ImportError:
            self.log_and_status("缺少必要的库: 请安装pywin32库")
        except Exception as e:
            self.log_and_status(f"转换过程中出错: {str(e)}")
        finally:
            # 确保WPS应用程序被正确关闭
            if wps_app:
                try:
                    wps_app.Quit()
                except:
                    pass
            
            # 清理COM资源
            try:
                pythoncom.CoUninitialize()
            except:
                pass

    def select_template_file(self):
        """
        选择模板文件
        """
        # 获取上次使用的模板目录
        last_template_dir = self.load_last_template_dir()
        
        file_path = filedialog.askopenfilename(
            title="选择模板文件",
            filetypes=[
                ("Word文档", "*.docx"),
                ("Excel文件", "*.xlsx"),
                ("所有文件", "*.*")
            ],
            initialdir=last_template_dir
        )
        
        if file_path:
            # 保存所选文件的目录
            file_dir = os.path.dirname(file_path)
            self.save_last_template_dir(file_dir)
            
            self.load_template_file(file_path)
        """
        加载Word模板文件
        :param file_path: Word文档路径
        """
        try:
            # 保存当前文件路径
            self.current_template_file = file_path
            self.current_file_label.config(text=f"当前文件: {os.path.basename(file_path)}")
            
            # 显示文档信息
            self.display_word_doc_info(file_path)
            
            print(f"已加载Word文档: {file_path}")
            print("点击'在Word中编辑'按钮开始编辑文档")
        except Exception as e:
            print(f"加载Word文档时出错: {str(e)}")

    def display_word_doc_info(self, file_path):
        """
        显示Word文档信息
        :param file_path: Word文档路径
        """
        try:
            doc = Document(file_path)
            
            # 收集文档信息
            info_lines = []
            info_lines.append(f"文件名: {os.path.basename(file_path)}")
            info_lines.append(f"文件路径: {file_path}")
            info_lines.append(f"段落数量: {len(doc.paragraphs)}")
            info_lines.append(f"表格数量: {len(doc.tables)}")
            
            # 显示前几个段落的内容预览
            info_lines.append("\n内容预览:")
            for i, paragraph in enumerate(doc.paragraphs[:10]):  # 只显示前10个段落
                if paragraph.text.strip():
                    info_lines.append(f"  段落 {i+1}: {paragraph.text[:50]}{'...' if len(paragraph.text) > 50 else ''}")
            
            if len(doc.paragraphs) > 10:
                info_lines.append(f"  ... 还有 {len(doc.paragraphs) - 10} 个段落")
            
            # 显示表格信息
            if doc.tables:
                info_lines.append(f"\n表格信息:")
                for i, table in enumerate(doc.tables):
                    info_lines.append(f"  表格 {i+1}: {len(table.rows)} 行 x {len(table.columns)} 列")
            
            # 更新文档信息显示
            self.doc_info_text.config(state=tk.NORMAL)
            self.doc_info_text.delete(1.0, tk.END)
            self.doc_info_text.insert(1.0, '\n'.join(info_lines))
            self.doc_info_text.config(state=tk.DISABLED)
            
            # 保存文档对象引用
            self.current_doc_object = doc
        except Exception as e:
            error_info = f"无法读取文档信息: {str(e)}"
            self.doc_info_text.config(state=tk.NORMAL)
            self.doc_info_text.delete(1.0, tk.END)
            self.doc_info_text.insert(1.0, error_info)
            self.doc_info_text.config(state=tk.DISABLED)

    def insert_placeholder_to_word(self):
        """
        将选中的占位符插入到Word/WPS文档的光标位置
        """
        # 检查是否选择了占位符
        selection = self.placeholder_listbox.curselection()
        if not selection:
            print("请先从左侧列表中选择一个占位符")
            return
        
        placeholder = self.placeholder_listbox.get(selection[0])
        
        # 检查是否已打开Word/WPS文档
        if not self.word_doc:
            print("请先点击'在Word中编辑'按钮打开Word或WPS文档")
            return
        
        try:
            import win32com.client
            
            # 在文档的光标位置插入占位符
            selection = self.word_app.Selection
            selection.TypeText(f"{{{placeholder}}}")
            
            print(f"已将占位符 {{{placeholder}}} 插入到文档")
        except Exception as e:
            print(f"插入占位符时出错: {str(e)}")
            print("请确保文档已正确打开")

    def save_word_template_file(self):
        """
        保存Word模板文件
        """
        if not hasattr(self, 'current_template_file') or not self.current_template_file:
            print("请先选择一个Word文档")
            return
        
        try:
            # 如果Word文档已打开，先保存
            if self.word_doc:
                self.word_doc.Save()
                print(f"已保存Word文档: {self.current_template_file}")
            else:
                print("文档已保存")
                
        except Exception as e:
            print(f"保存Word文档时出错: {str(e)}")

    def refresh_placeholders_list(self):
        """
        刷新占位符列表
        """
        # 在另一个线程中执行占位符刷新操作
        import threading
        threading.Thread(target=self._refresh_placeholders_thread, daemon=True).start()

    def edit_in_word(self):
        """
        在Word中编辑当前文档
        """
        if not hasattr(self, 'current_template_file') or not self.current_template_file:
            print("请先选择一个Word文档")
            return
        
        if not self.current_template_file.endswith('.docx'):
            print("只能在Word中编辑.docx文件")
            return
        
        try:
            import win32com.client
            import os
            
            # 获取文件的绝对路径
            abs_path = os.path.abspath(self.current_template_file)
            
            # 启动Word应用程序
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = True  # 显示Word应用程序
            
            # 打开文档
            doc = word.Documents.Open(abs_path)
            
            print(f"已在Word中打开文档: {self.current_template_file}")
            print("请在Word中进行编辑，关闭文档时会自动保存更改")
            
        except Exception as e:
            print(f"在Word中编辑文档时出错: {str(e)}")
            print("请确保：")
            print("1. 已安装Microsoft Word")
            print("2. 已安装pywin32库 (pip install pywin32)")
            print("3. 以足够权限运行程序")
    
    
    def _refresh_placeholders_thread(self):
        """
        在线程中执行占位符刷新操作
        """
        # 更新状态栏
        # self.update_status("正在刷新占位符列表...")
        
        # 在主线程中更新UI
        self.root.after(0, lambda: self.placeholder_listbox.delete(0, tk.END))
        
        # 检查用户是否已选择模板目录
        if not hasattr(self, 'selected_template_folder') or not self.selected_template_folder:
            # 在主线程中更新UI
            self.root.after(0, lambda: [
                self.placeholder_listbox.delete(0, tk.END),
                self.placeholder_listbox.insert(tk.END, "日期"),
                self.placeholder_listbox.insert(tk.END, "请先点击\"选择文档目录\"按钮选择文件夹"),
                self.placeholder_listbox.itemconfig(0, {'fg': 'gray'}),
                self.placeholder_listbox.itemconfig(1, {'fg': 'gray'}),
                self.delete_placeholder_button.config(state=tk.DISABLED)  # 禁用删除按钮
            ])
            # self.root.after(0, lambda: self.update_status("请先选择模板目录"))
            return
        
        # 从用户选择的模板目录中收集所有占位符
        all_placeholders = set()
        placeholder_files = {}  # 记录每个占位符出现在哪些文件中
        try:
            # 遍历文件夹中的所有文件
            for file in os.listdir(self.selected_template_folder):
                file_path = os.path.join(self.selected_template_folder, file)
                if os.path.isfile(file_path):
                    # 根据文件扩展名处理不同类型的文件
                    if file.endswith('.docx'):
                        placeholders = self.processor.extract_placeholders_from_docx(file_path)
                        all_placeholders.update(placeholders)
                    elif file.endswith('.xlsx'):
                        placeholders = self.processor.extract_placeholders_from_xlsx(file_path)
                        all_placeholders.update(placeholders)
                    else:
                        continue
                    
                    # 记录每个占位符出现的文件
                    for placeholder in placeholders:
                        if placeholder not in placeholder_files:
                            placeholder_files[placeholder] = []
                        placeholder_files[placeholder].append(file_path)
            
            # 保存占位符和文件的映射关系
            self.placeholder_files = placeholder_files
            
            # 在主线程中更新UI
            def update_ui():
                # 清空占位符列表
                self.placeholder_listbox.delete(0, tk.END)
                
                # 添加占位符到列表
                if not all_placeholders:
                    self.placeholder_listbox.insert(tk.END, "日期")
                    self.placeholder_listbox.insert(tk.END, "在选定的文件夹中未找到占位符")
                    self.placeholder_listbox.itemconfig(0, {'fg': 'black'})  # 用户已选择目录，日期占位符变为可用
                    self.placeholder_listbox.itemconfig(1, {'fg': 'gray'})
                    # 禁用删除按钮
                    self.delete_placeholder_button.config(state=tk.DISABLED)
                    # self.update_status("在选定的文件夹中未找到占位符")
                else:
                    # 添加日期占位符（始终在列表顶部，但避免重复）
                    self.placeholder_listbox.insert(tk.END, "日期")
                    self.placeholder_listbox.itemconfig(0, {'fg': 'black'})  # 用户已选择目录，日期占位符变为可用
                    
                    # 添加其他占位符到列表（排除已存在的日期占位符）
                    for placeholder in sorted(all_placeholders):
                        if placeholder != "日期":  # 避免重复添加日期占位符
                            self.placeholder_listbox.insert(tk.END, placeholder)
                    
                    # 启用删除按钮
                    self.delete_placeholder_button.config(state=tk.NORMAL)
                    # self.update_status(f"刷新完成，找到 {len(all_placeholders)} 个占位符")
            
            self.root.after(0, update_ui)
            
        except Exception as e:
            # 在主线程中更新UI
            self.root.after(0, lambda: [
                self.placeholder_listbox.delete(0, tk.END),
                self.placeholder_listbox.insert(tk.END, "日期"),
                self.placeholder_listbox.insert(tk.END, f"读取占位符出错: {str(e)}"),
                self.placeholder_listbox.itemconfig(0, {'fg': 'black'}),  # 即使出错，日期占位符也应该是可用的
                self.placeholder_listbox.itemconfig(1, {'fg': 'red'}),
                self.update_status(f"刷新占位符时出错: {str(e)}"),
                self.delete_placeholder_button.config(state=tk.DISABLED)  # 禁用删除按钮
            ])
    
    def refresh_placeholders(self):
        """
        刷新占位符按钮的处理函数
        """
        # 直接调用refresh_placeholders_list方法来重新获取并显示占位符
        self.refresh_placeholders_list()
    
    def edit_placeholder(self, event=None):
        """
        双击编辑占位符
        """
        # 检查是否选择了占位符
        selection = self.placeholder_listbox.curselection()
        if not selection:
            self.log_and_status("请先从列表中选择一个占位符")
            return
        
        # 获取当前选中的占位符
        index = selection[0]
        old_placeholder = self.placeholder_listbox.get(index)
        
        # 检查是否为提示信息
        if old_placeholder in ['''请先点击"选择文档目录"按钮选择文件夹''', 
                              "在选定的文件夹中未找到占位符",
                              "日期"]:
            self.log_and_status("无法编辑此内容")
            return
        
        # 创建编辑对话框
        dialog = tk.Toplevel(self.root)
        
        # 设置对话框大小并居中
        dialog.geometry("250x100")
        dialog.resizable(False, False)
        
        # 居中显示对话框
        self.center_dialog(dialog, 250, 100)
        dialog.transient(self.root)
        dialog.grab_set()

        # 设置标题
        dialog.title("编辑占位符")
      
        # 设置对话框图标
        self.set_dialog_icon(dialog)
        
        # 编辑框架，设置整体居中
        edit_frame = ttk.Frame(dialog)
        edit_frame.grid(row=0, column=0, pady=5, padx=5, sticky=(tk.W, tk.E))
        
        # 配置编辑框架的列权重
        edit_frame.columnconfigure(0, weight=1)
        edit_frame.columnconfigure(1, weight=2)
        
        # 添加标签和输入框
        ttk.Label(edit_frame, text="占位符名称:").grid(row=0, column=0, pady=(5, 5), sticky=tk.W)
        
        placeholder_var = tk.StringVar(value=old_placeholder)
        entry = ttk.Entry(edit_frame, textvariable=placeholder_var)
        entry.grid(row=0, column=1, pady=5, padx=(5, 5), sticky=(tk.W, tk.E))
        entry.select_range(0, tk.END)
        entry.focus()
        
        # 确定按钮事件处理
        def on_ok():
            new_placeholder = placeholder_var.get().strip()
            if not new_placeholder:
                self.log_and_status("占位符名称不能为空")
                return
            
            if new_placeholder == old_placeholder:
                dialog.destroy()
                return
            
            # 在另一个线程中更新占位符，防止阻塞主界面
            def update_thread():
                try:
                    self.update_placeholder_in_templates(old_placeholder, new_placeholder)
                    # 更新列表框中的显示
                    self.placeholder_listbox.delete(index)
                    self.placeholder_listbox.insert(index, new_placeholder)
                    self.log_and_status(f"占位符已从 '{old_placeholder}' 更新为 '{new_placeholder}'")
                    dialog.destroy()
                except Exception as e:
                    self.log_and_status(f"更新占位符时出错: {str(e)}")
            
            # 启动线程执行更新操作
            import threading
            thread = threading.Thread(target=update_thread, daemon=True)
            thread.start()
        
        # 取消按钮事件处理
        def on_cancel():
            dialog.destroy()
        
        # 按钮框架，设置整体居中
        button_frame = ttk.Frame(dialog)
        button_frame.grid(row=1, column=0, pady=5, padx=5, sticky=(tk.W, tk.E))
        
        # 配置按钮框架的列权重，使按钮能够居中
        button_frame.columnconfigure(0, weight=1)
        button_frame.columnconfigure(1, weight=0)
        button_frame.columnconfigure(2, weight=0)
        button_frame.columnconfigure(3, weight=1)
        
        # 确认按钮 取消按钮
        ok_button = ttk.Button(button_frame, text="确定", command=on_ok)
        ok_button.grid(row=0, column=1, padx=(0, 5))
        
        cancel_button = ttk.Button(button_frame, text="取消", command=on_cancel)
        cancel_button.grid(row=0, column=2, padx=(5, 0))
        
        # 配置对话框的回车和ESC键事件
        entry.bind('<Return>', lambda e: on_ok())
        dialog.bind('<Escape>', lambda e: on_cancel())
    
    def update_placeholder_in_templates(self, old_placeholder, new_placeholder):
        """
        在模板中更新占位符
        :param old_placeholder: 旧占位符
        :param new_placeholder: 新占位符
        """
        # 检查用户是否已选择模板目录
        if not hasattr(self, 'selected_template_folder') or not self.selected_template_folder:
            raise Exception("未选择模板目录")
        
        # 遍历模板目录中的所有.docx和.xlsx文件
        updated_files = []
        for file in os.listdir(self.selected_template_folder):
            file_path = os.path.join(self.selected_template_folder, file)
            if os.path.isfile(file_path):
                try:
                    if file.endswith('.docx'):
                        self.update_placeholder_in_docx(file_path, old_placeholder, new_placeholder)
                        updated_files.append(file)
                    elif file.endswith('.xlsx'):
                        self.update_placeholder_in_xlsx(file_path, old_placeholder, new_placeholder)
                        updated_files.append(file)
                except Exception as e:
                    print(f"更新文件 {file} 中的占位符时出错: {e}")
        
        if not updated_files:
            print("未找到需要更新的模板文件")
        else:
            print(f"已在以下文件中更新占位符: {', '.join(updated_files)}")
    
    def update_placeholder_in_docx(self, file_path, old_placeholder, new_placeholder):
        """
        在Word文档中更新占位符
        :param file_path: Word文档路径
        :param old_placeholder: 旧占位符名称
        :param new_placeholder: 新占位符名称
        """
        from docx import Document
        
        # 打开文档
        doc = Document(file_path)
        
        # 替换段落中的占位符
        for paragraph in doc.paragraphs:
            if f"{{{old_placeholder}}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"{{{old_placeholder}}}", f"{{{new_placeholder}}}")
        
        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if f"{{{old_placeholder}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{old_placeholder}}}", f"{{{new_placeholder}}}")
        
        # 保存文档
        doc.save(file_path)
    
    def update_placeholder_in_xlsx(self, file_path, old_placeholder, new_placeholder):
        """
        在Excel文件中更新占位符
        :param file_path: Excel文件路径
        :param old_placeholder: 旧占位符名称
        :param new_placeholder: 新占位符名称
        """
        from openpyxl import load_workbook
        
        # 打开工作簿
        workbook = load_workbook(file_path)
        
        # 遍历所有工作表
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            
            # 遍历所有单元格
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str) and f"{{{old_placeholder}}}" in cell.value:
                        cell.value = cell.value.replace(f"{{{old_placeholder}}}", f"{{{new_placeholder}}}")
        
        # 保存工作簿
        workbook.save(file_path)
    
    def add_new_placeholder(self):
        """
        添加新占位符
        """
        # 创建添加占位符对话框
        dialog = tk.Toplevel(self.root)
        
        # 设置对话框大小并居中
        dialog.geometry("250x100")
        dialog.resizable(False, False)
        dialog.title("添加新占位符")
        
        # 居中显示对话框
        self.center_dialog(dialog, 250, 100)
        dialog.transient(self.root)
        dialog.grab_set()
        
        # 设置对话框图标
        self.set_dialog_icon(dialog)
        
        # 创建录入框架，包含标签和输入框，整体居中
        input_frame = ttk.Frame(dialog)
        input_frame.grid(row=0, column=0, pady=5, padx=5, sticky=(tk.W, tk.E))
        
        # 配置录入框架的列权重
        input_frame.columnconfigure(0, weight=1)
        input_frame.columnconfigure(1, weight=2)
        
        # 添加标签和输入框
        ttk.Label(input_frame, text="占位符名称:").grid(row=0, column=0, pady=5, sticky=tk.W)
        
        placeholder_var = tk.StringVar()
        entry = ttk.Entry(input_frame, textvariable=placeholder_var)
        entry.grid(row=0, column=1, pady=5, padx=5, sticky=(tk.W, tk.E))
        entry.focus()
        
        # 确定按钮事件处理
        def on_ok():
            new_placeholder = placeholder_var.get().strip()
            if not new_placeholder:
                self.log_and_status("占位符名称不能为空")
                return
            
            # 检查占位符是否已存在
            for i in range(self.placeholder_listbox.size()):
                if self.placeholder_listbox.get(i) == new_placeholder:
                    self.log_and_status(f"占位符 '{new_placeholder}' 已存在")
                    return
            
            # 添加占位符到列表
            self.placeholder_listbox.insert(tk.END, new_placeholder)
            self.log_and_status(f"已添加新占位符: {new_placeholder}")
            
            # 如果用户已选择模板目录，则在模板中添加占位符
            if hasattr(self, 'selected_template_folder') and self.selected_template_folder:
                self.add_placeholder_to_templates(new_placeholder)
            
            dialog.destroy()
        
        # 取消按钮事件处理
        def on_cancel():
            dialog.destroy()
        
        # 按钮框架
        button_frame = ttk.Frame(dialog)
        button_frame.grid(row=1, column=0, pady=5, padx=5, sticky=(tk.W, tk.E))

        # 配置按钮框架的列权重，使按钮能够居中
        button_frame.columnconfigure(0, weight=1)
        button_frame.columnconfigure(1, weight=0)
        button_frame.columnconfigure(2, weight=0)
        button_frame.columnconfigure(3, weight=1)
        
        # 确定和取消按钮
        ok_button = ttk.Button(button_frame, text="确定", command=on_ok)
        ok_button.grid(row=0, column=1, padx=(0, 5))
        cancel_button = ttk.Button(button_frame, text="取消", command=on_cancel)
        cancel_button.grid(row=0, column=2, padx=(5, 0))
        
        # 绑定回车键到确定按钮
        entry.bind('<Return>', lambda e: on_ok())
        # 绑定ESC键到取消按钮
        dialog.bind('<Escape>', lambda e: on_cancel())

    def delete_placeholder(self):
        """
        删除选中的占位符
        """
        # 检查是否选择了占位符
        selection = self.placeholder_listbox.curselection()
        if not selection:
            self.log_and_status("请先从列表中选择一个占位符")
            return
        
        # 获取当前选中的占位符
        index = selection[0]
        old_placeholder = self.placeholder_listbox.get(index)
        
        # 检查是否为日期占位符
        if old_placeholder == "日期":
            self.log_and_status("无法编辑默认的日期占位符")
            return
        
        # 检查是否为提示信息
        if old_placeholder in ["请先点击\"选择文档目录\"按钮选择文件夹", 
                              "在选定的文件夹中未找到占位符"]:
            self.log_and_status("无法编辑提示信息")
            return
        
        # 确认删除操作
        from tkinter import messagebox
        result = messagebox.askyesno("确认删除", f"确定要删除占位符 '{old_placeholder}' 吗？\n这将从所有模板文件中移除此占位符。")
        if not result:
            return
        
        try:
            # 从模板文件中删除占位符（用空字符串替换）
            self.remove_placeholder_from_templates(old_placeholder)
            
            # 从列表框中删除占位符
            self.placeholder_listbox.delete(index)
            
            self.log_and_status(f"已删除占位符: {old_placeholder}")
        except Exception as e:
            self.log_and_status(f"删除占位符时出错: {str(e)}")

    def remove_placeholder_from_templates(self, placeholder):
        """
        从模板中删除占位符
        :param placeholder: 要删除的占位符
        """
        # 检查用户是否已选择模板目录
        if not hasattr(self, 'selected_template_folder') or not self.selected_template_folder:
            raise Exception("未选择模板目录")
        
        # 遍历模板目录中的所有.docx和.xlsx文件
        updated_files = []
        for file in os.listdir(self.selected_template_folder):
            file_path = os.path.join(self.selected_template_folder, file)
            if os.path.isfile(file_path):
                try:
                    if file.endswith('.docx'):
                        self.remove_placeholder_from_docx(file_path, placeholder)
                        updated_files.append(file)
                    elif file.endswith('.xlsx'):
                        self.remove_placeholder_from_xlsx(file_path, placeholder)
                        updated_files.append(file)
                except Exception as e:
                    print(f"更新文件 {file} 中的占位符时出错: {e}")
        
        if not updated_files:
            print("未找到需要更新的模板文件")
        else:
            print(f"已在以下文件中删除占位符: {', '.join(updated_files)}")

    def remove_placeholder_from_docx(self, file_path, placeholder):
        """
        从Word文档中删除占位符
        :param file_path: Word文档路径
        :param placeholder: 要删除的占位符名称
        """
        from docx import Document
        
        # 打开文档
        doc = Document(file_path)
        
        # 替换段落中的占位符为空字符串
        for paragraph in doc.paragraphs:
            if f"{{{placeholder}}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"{{{placeholder}}}", "")
        
        # 替换表格中的占位符为空字符串
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if f"{{{placeholder}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{placeholder}}}", "")
        
        # 保存文档
        doc.save(file_path)

    def remove_placeholder_from_xlsx(self, file_path, placeholder):
        """
        从Excel文件中删除占位符
        :param file_path: Excel文件路径
        :param placeholder: 要删除的占位符名称
        """
        from openpyxl import load_workbook
        
        # 打开工作簿
        workbook = load_workbook(file_path)
        
        # 遍历所有工作表
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            
            # 遍历所有单元格
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str) and f"{{{placeholder}}}" in cell.value:
                        cell.value = cell.value.replace(f"{{{placeholder}}}", "")
        
        # 保存工作簿
        workbook.save(file_path)

    def add_placeholder_to_templates(self, placeholder):
        """
        在模板中添加占位符（这是一个空实现，因为占位符是在文档中手动添加的）
        :param placeholder: 要添加的占位符
        """
        # 实际上，占位符是在文档中手动添加的，这里不需要做任何事情
        pass
    
    def add_placeholder_button_frame(self):
        """
        添加添加占位符按钮
        """
        add_placeholder_button_frame = ttk.Frame(self.root)
        add_placeholder_button_frame.pack(side=tk.TOP, fill=tk.X, padx=10, pady=10)
        
        ttk.Button(add_placeholder_button_frame, text="添加占位符", command=self.add_placeholder_dialog).pack(side=tk.LEFT)
    
    def copy_placeholder_to_clipboard(self):
        """
        将选中的占位符复制到剪贴板
        """
        selection = self.placeholder_listbox.curselection()
        if not selection:
            print("请先选择一个占位符")
            return
        
        placeholder = self.placeholder_listbox.get(selection[0])
        formatted_placeholder = f"{{{placeholder}}}"
        # 使用Tkinter内置剪贴板功能替代pyperclip
        try:
            self.root.clipboard_clear()
            self.root.clipboard_append(formatted_placeholder)
            self.root.update()
            self.update_status(f"已复制占位符: {formatted_placeholder}，请在文档中粘贴")
        except Exception as e:
            self.update_status(f"复制失败: {e}")
    
    def insert_placeholder_at_cursor(self):
        """
        将选中的占位符添加到剪贴板（原为在光标位置插入占位符）
        """
        selection = self.placeholder_listbox.curselection()
        if not selection:
            print("请先选择一个占位符")
            return
        
        placeholder = self.placeholder_listbox.get(selection[0])
        # 复制占位符到剪贴板
        self.copy_placeholder_to_clipboard()
    
    def select_template_file(self):
        """
        选择模板文件
        """
        file_path = filedialog.askopenfilename(
            title="选择模板文件",
            filetypes=[
                ("Word文档", "*.docx"),
                ("Excel文件", "*.xlsx"),
                ("所有文件", "*.*")
            ]
        )
        
        if file_path:
            self.load_template_file(file_path)

    def load_template_file(self, file_path):
        """
        加载模板文件内容
        :param file_path: 文件路径
        """
        try:
            # 根据文件类型处理
            if file_path.endswith('.docx'):
                self.load_docx_content(file_path)
            elif file_path.endswith('.xlsx'):
                self.load_xlsx_content(file_path)
            else:
                # 默认按文本文件处理，显示基本信息
                info_lines = []
                info_lines.append(f"文件名: {os.path.basename(file_path)}")
                info_lines.append(f"文件路径: {file_path}")
                info_lines.append("注意: 这是非Word/Excel文档格式")
                info_lines.append("仅支持占位符复制功能")
                
                # 更新文档信息显示
                self.doc_info_text.config(state=tk.NORMAL)
                self.doc_info_text.delete(1.0, tk.END)
                self.doc_info_text.insert(1.0, '\n'.join(info_lines))
                self.doc_info_text.config(state=tk.DISABLED)
                
                self.current_file_label.config(text=f"当前文件: {os.path.basename(file_path)}")
                self.current_template_file = file_path
                print(f"已加载文件: {file_path}")
                print("这是非Word文档格式，仅支持占位符复制功能")
            
            self.current_file_label.config(text=f"当前文件: {os.path.basename(file_path)}")
            self.current_template_file = file_path
            
            # 自动打开文档（无论什么格式）
            self.auto_open_document(file_path)
        except Exception as e:
            error_info = f"无法读取文档信息: {str(e)}"
            self.doc_info_text.config(state=tk.NORMAL)
            self.doc_info_text.delete(1.0, tk.END)
            self.doc_info_text.insert(1.0, error_info)
            self.doc_info_text.config(state=tk.DISABLED)
            print(error_info)
    
    def add_placeholder_button_frame(self):
        """
        添加添加占位符按钮
        """
        add_placeholder_button_frame = ttk.Frame(self.root)
        add_placeholder_button_frame.pack(side=tk.TOP, fill=tk.X, padx=10, pady=10)
        
        ttk.Button(add_placeholder_button_frame, text="添加占位符", command=self.add_placeholder_dialog).pack(side=tk.LEFT)
        ttk.Button(add_placeholder_button_frame, text="复制占位符到剪贴板", command=self.copy_placeholder_to_clipboard).pack()
        
    def on_placeholder_select(self, event=None):
        """
        当用户在占位符列表中选择一个占位符时，显示使用该占位符的文件
        """
        selection = self.placeholder_listbox.curselection()
        if not selection:
            return
            
        # 获取选中的占位符
        placeholder = self.placeholder_listbox.get(selection[0])
        
        # 检查是否为提示信息
        if placeholder in ["请先点击\"打开模板目录\"按钮选择文件夹", 
                          "在选定的文件夹中未找到占位符"]:
            return
            
        # 显示使用该占位符的文件
        self.show_files_for_placeholder(placeholder)
        
    def show_files_for_placeholder(self, placeholder):
        """
        显示使用指定占位符的文件列表
        :param placeholder: 占位符名称
        """
        # 检查是否有该占位符的文件映射信息
        if placeholder not in self.placeholder_files:
            info_text = f"未找到使用占位符 {{{placeholder}}} 的文件信息"
            self.doc_info_text.config(state=tk.NORMAL)
            self.doc_info_text.delete(1.0, tk.END)
            self.doc_info_text.insert(1.0, info_text)
            self.doc_info_text.config(state=tk.DISABLED)
            return
            
        # 获取使用该占位符的文件列表
        files = self.placeholder_files[placeholder]
        
        # 构造显示文本
        info_lines = [f"使用占位符 {{{placeholder}}} 的文件:"]
        info_lines.append("-" * 40)
        
        for i, file_path in enumerate(files, 1):
            file_name = os.path.basename(file_path)
            info_lines.append(f"{i}. {file_name}")
            
        info_lines.append("-" * 40)
        info_lines.append(f"共 {len(files)} 个文件使用此占位符")
        
        # 更新文档信息显示区域
        self.doc_info_text.config(state=tk.NORMAL)
        self.doc_info_text.delete(1.0, tk.END)
        self.doc_info_text.insert(1.0, "\n".join(info_lines))
        self.doc_info_text.config(state=tk.DISABLED)

    def auto_open_document(self, file_path):
        """
        自动打开文档
        :param file_path: 文件路径
        """
        try:
            import win32com.client
            import os
            
            # 获取文件的绝对路径
            abs_path = os.path.abspath(file_path)
            
            # 根据文件扩展名确定应用程序
            if file_path.endswith('.docx'):
                # 尝试连接到正在运行的Word实例
                try:
                    word = win32com.client.GetActiveObject("Word.Application")
                except:
                    # 如果没有运行的Word实例，则启动新的实例
                    word = win32com.client.Dispatch("Word.Application")
                
                word.Visible = True
                doc = word.Documents.Open(abs_path)
                print(f"已在Word中打开文档: {file_path}")
                
            elif file_path.endswith('.xlsx'):
                # 尝试连接到正在运行的Excel实例
                try:
                    excel = win32com.client.GetActiveObject("Excel.Application")
                except:
                    # 如果没有运行的Excel实例，则启动新的实例
                    excel = win32com.client.Dispatch("Excel.Application")
                
                excel.Visible = True
                workbook = excel.Workbooks.Open(abs_path)
                print(f"已在Excel中打开文档: {file_path}")
                
            else:
                # 对于其他格式的文件，使用系统默认程序打开
                try:
                    import subprocess
                    if os.name == 'nt':  # Windows系统
                        os.startfile(abs_path)
                    elif os.name == 'posix':  # macOS或Linux系统
                        subprocess.call(['open', abs_path])  # macOS
                except:
                    try:
                        subprocess.call(['xdg-open', abs_path])  # Linux
                    except:
                        print(f"无法自动打开文件: {file_path}")
                        print("请手动打开该文件进行编辑")
                
                print(f"已尝试使用默认程序打开文档: {file_path}")
                
        except Exception as e:
            # 如果自动打开失败，只打印错误信息，不中断程序流程
            print(f"自动打开文档时出错: {str(e)}")
            print("请手动打开文档进行编辑")

    def load_docx_content(self, file_path):
        """
        加载Word文档内容（仅显示信息，不显示内容）
        :param file_path: Word文档路径
        """
        try:
            # 显示文档信息
            self.display_word_doc_info(file_path)
        except Exception as e:
            print(f"加载Word文档时出错: {str(e)}")
    
    def load_xlsx_content(self, file_path):
        """
        加载Excel文件内容（仅显示信息，不显示内容）
        :param file_path: Excel文件路径
        """
        try:
            if not EXCEL_PROCESSING_AVAILABLE:
                raise Exception("Excel处理功能不可用，请安装openpyxl库")
            
            workbook = load_workbook(file_path)
            
            # 收集文档信息
            info_lines = []
            info_lines.append(f"文件名: {os.path.basename(file_path)}")
            info_lines.append(f"文件路径: {file_path}")
            info_lines.append(f"工作表数量: {len(workbook.sheetnames)}")
            
            # 显示工作表信息
            info_lines.append("\n工作表列表:")
            for i, sheet_name in enumerate(workbook.sheetnames):
                worksheet = workbook[sheet_name]
                info_lines.append(f"  {i+1}. {sheet_name} ({worksheet.max_row} 行 x {worksheet.max_column} 列)")
            
            # 更新文档信息显示
            self.doc_info_text.config(state=tk.NORMAL)
            self.doc_info_text.delete(1.0, tk.END)
            self.doc_info_text.insert(1.0, '\n'.join(info_lines))
            self.doc_info_text.config(state=tk.DISABLED)
            
            print(f"已加载Excel文件: {file_path}")
        except Exception as e:
            error_info = f"无法读取文档信息: {str(e)}"
            self.doc_info_text.config(state=tk.NORMAL)
            self.doc_info_text.delete(1.0, tk.END)
            self.doc_info_text.insert(1.0, error_info)
            self.doc_info_text.config(state=tk.DISABLED)
            print(f"加载Excel文件时出错: {str(e)}")
    
    def setup_main_tab(self):
        """
        设置主操作标签页
        """
        # 主框架
        main_frame = ttk.Frame(self.main_frame, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 配置网格权重
        main_frame.columnconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(0, weight=1)
        main_frame.rowconfigure(1, weight=0)
        
        # 创建左右两栏框架
        left_frame = ttk.LabelFrame(main_frame, text="方案选择", padding="10")
        left_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(0, 5))
        
        right_frame = ttk.LabelFrame(main_frame, text="用户录入", padding="10")
        right_frame.grid(row=0, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(5, 0))
        
        # 左侧：方案列表
        self.scheme_listbox_main = tk.Listbox(left_frame, height=15)
        self.scheme_listbox_main.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))
        self.scheme_listbox_main.bind('<<ListboxSelect>>', self.on_scheme_select)
        
        left_frame.columnconfigure(0, weight=1)
        left_frame.columnconfigure(1, weight=1)
        left_frame.rowconfigure(0, weight=1)
        
        ttk.Button(left_frame, text="刷新方案列表", command=self.load_saved_schemes).grid(row=1, column=0, columnspan=2)
        
        # 右侧：用户录入区域
        self.input_canvas = tk.Canvas(right_frame, height=200, relief=tk.SUNKEN, borderwidth=1)
        self.input_scrollbar = ttk.Scrollbar(right_frame, orient="vertical", command=self.input_canvas.yview)
        self.input_scrollable_frame = ttk.Frame(self.input_canvas)
        
        # 添加鼠标滚轮支持
        def _on_mousewheel(event):
            self.input_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        
        self.input_canvas.bind("<MouseWheel>", _on_mousewheel)
        self.input_scrollable_frame.bind("<MouseWheel>", _on_mousewheel)
        
        self.input_scrollable_frame.bind(
            "<Configure>",
            lambda e: self.input_canvas.configure(
                scrollregion=self.input_canvas.bbox("all")
            )
        )
        
        self.input_canvas.create_window((0, 0), window=self.input_scrollable_frame, anchor="nw")
        self.input_canvas.configure(yscrollcommand=self.input_scrollbar.set)
        
        self.input_canvas.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.input_scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        
        # 添加保存按钮框架
        save_button_frame = ttk.Frame(right_frame)
        save_button_frame.grid(row=1, column=0, columnspan=2, pady=(5, 0), sticky=(tk.W, tk.E))
        
        # 修改为左右布局
        # Label和Combobox作为一个整体居左放置，按钮居中放置
        # 创建包含Label和Combobox的子框架，作为一个整体
        history_frame = ttk.Frame(save_button_frame)
        history_frame.grid(row=0, column=0, padx=(0, 5), pady=5, sticky=tk.W)
        
        ttk.Label(history_frame, text="历史记录:").pack(side=tk.LEFT)
        self.history_combobox = ttk.Combobox(history_frame, width=20, state="disabled")
        self.history_combobox.pack(side=tk.LEFT, padx=(5, 0))
        # 绑定选择事件
        self.history_combobox.bind("<<ComboboxSelected>>", self.load_history_record)
        
        # 添加删除历史记录按钮
        self.delete_history_button = ttk.Button(history_frame, text="删除记录", width=8, command=self.delete_history_record, state="disabled")
        self.delete_history_button.pack(side=tk.LEFT, padx=(5, 0))
        
        # 保存按钮居中放置
        self.save_inputs_button = ttk.Button(save_button_frame, text="保存录入内容", command=self.save_user_inputs, state="disabled")
        self.save_inputs_button.grid(row=0, column=1, pady=5)

        # 配置保存按钮框架的列权重，使按钮能够居中
        save_button_frame.columnconfigure(0, weight=1)  # 左侧组件
        save_button_frame.columnconfigure(1, weight=0)  # 中间按钮（不拉伸）
        save_button_frame.columnconfigure(2, weight=1)  # 右侧空白区域（用于居中按钮）
        
        right_frame.columnconfigure(0, weight=1)
        right_frame.rowconfigure(0, weight=1)
        
        # 控制按钮区域
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=1, column=0, columnspan=2, pady=(10, 0))
        
        # 输出目录选择按钮
        ttk.Button(button_frame, text="选择输出文件夹", command=self.select_output_dir).grid(row=0, column=0, padx=(0, 10))
        self.output_dir_label = ttk.Label(button_frame, text=f"输出目录: {self.output_dir}", width=45, anchor="w")
        self.output_dir_label.grid(row=0, column=1, padx=(0, 10))
        
        self.generate_docs_button = ttk.Button(button_frame, text="生成文档", command=self.generate_documents, state="disabled")
        self.generate_docs_button.grid(row=0, column=2, padx=(10, 10))
        
        # 如果PDF功能可用，添加合并为PDF按钮（在打开输出文件夹按钮之前）
        if PDF_CONVERSION_AVAILABLE and PDF_MERGING_AVAILABLE:
            self.merge_pdf_button = ttk.Button(button_frame, text="合并为PDF", command=self.merge_to_pdf, state="disabled")
            self.merge_pdf_button.grid(row=0, column=3, padx=(10, 10))
            ttk.Button(button_frame, text="打开输出文件夹", command=self.open_output_dir).grid(row=0, column=4, padx=(10, 10))
        elif not PDF_CONVERSION_AVAILABLE or not PDF_MERGING_AVAILABLE:
            ttk.Button(button_frame, text="合并为PDF(需要安装依赖)", state=tk.DISABLED).grid(row=0, column=3, padx=(10, 10))
            ttk.Button(button_frame, text="打开输出文件夹", command=self.open_output_dir).grid(row=0, column=4, padx=(10, 10))
        else:
            ttk.Button(button_frame, text="打开输出文件夹", command=self.open_output_dir).grid(row=0, column=3, padx=(10, 10))
        
        # 加载已保存的方案
        self.load_saved_schemes()
        
        # 初始化时禁用相关控件
        self.disable_scheme_related_controls()
    
    def delete_history_record(self):
        """
        删除选中的历史记录
        """
        try:
            # 获取选中项的索引
            selected_index = self.history_combobox.current()
            if selected_index < 0:
                self.log_and_status("请先选择一条历史记录")
                return
            
            # 确认删除操作
            from tkinter import messagebox
            result = messagebox.askyesno("确认删除", "确定要删除选中的历史记录吗？此操作不可恢复。")
            if not result:
                return
            
            # 读取历史记录
            if not os.path.exists("app_data.json"):
                return
            
            with open("app_data.json", "r", encoding="utf-8") as f:
                data = json.load(f)
            
            history_data = data.get("history", {})
            
            # 获取当前方案的历史记录
            if self.current_scheme not in history_data:
                return
            
            # 删除选中的记录
            if selected_index < len(history_data[self.current_scheme]):
                del history_data[self.current_scheme][selected_index]
                
                # 保存更新后的历史记录
                data["history"] = history_data
                with open("app_data.json", "w", encoding="utf-8") as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
                
                # 更新下拉框内容
                self.update_history_combobox()
                
                # 清空选择
                self.history_combobox.set('')
                
                self.log_and_status("历史记录已删除")
            else:
                self.log_and_status("未找到选中的历史记录")
                
        except Exception as e:
            self.log_and_status(f"删除历史记录时出错: {e}")

    def select_output_dir(self):
        """
        选择输出文件夹
        """
        selected_dir = filedialog.askdirectory(
            title="选择输出文件夹",
            initialdir=self.output_dir
        )
        
        if selected_dir:
            self.output_dir = selected_dir
            self.save_last_output_dir(selected_dir)
            # 限制输出目录显示长度，防止界面布局混乱
            display_dir = self.output_dir
            if len(display_dir) > 60:
                display_dir = display_dir[:57] + "..."
            self.output_dir_label.config(text=f"输出目录: {display_dir}")

    def on_scheme_select(self, event):
        """
        当用户选择方案时，加载该方案的占位符和历史记录
        """
        selection = self.scheme_listbox_main.curselection()
        if not selection:
            return

        # 获取选中的方案名称
        selected_scheme = self.scheme_listbox_main.get(selection[0])
        self.current_scheme = selected_scheme

        try:
            # 读取方案数据
            with open("app_data.json", "r", encoding="utf-8") as f:
                data = json.load(f)

            schemes = data.get("schemes", {})
            if selected_scheme not in schemes:
                self.log_and_status(f"错误: 方案 '{selected_scheme}' 不存在")
                return

            scheme_data = schemes[selected_scheme]

            # 更新模板文件列表
            self.template_files = scheme_data.get("template_files", [])

            # 更新占位符列表
            self.ordered_placeholders = scheme_data.get("placeholder_order", [])
            
            # 清除现有的输入字段
            for widget in self.input_scrollable_frame.winfo_children():
                widget.destroy()
            
            self.input_fields = {}  # 重置输入字段字典
            
            # 创建新的输入字段
            self.create_input_fields()
            
            # 更新历史记录下拉框
            self.update_history_combobox()
            
            # 启用相关按钮和控件
            self.enable_scheme_related_controls()
            
            self.log_and_status(f"已加载方案: {selected_scheme}")
            
        except Exception as e:
            self.log_and_status(f"加载方案时出错: {str(e)}")
    
    def enable_scheme_related_controls(self):
        """
        启用与方案相关的控件
        """
        # 启用历史记录下拉框
        self.history_combobox.config(state="readonly")
        
        # 启用删除记录按钮
        self.delete_history_button.config(state="normal")
        
        # 启用保存录入内容按钮
        self.save_inputs_button.config(state="normal")
        
        # 启用生成文档和合并为PDF按钮
        self.generate_docs_button.config(state="normal")
        
        if hasattr(self, 'merge_pdf_button'):
            self.merge_pdf_button.config(state="normal")
    
    def disable_scheme_related_controls(self):
        """
        禁用与方案相关的控件
        """
        # 禁用历史记录下拉框
        self.history_combobox.config(state="disabled")
        
        # 禁用保存录入内容按钮（需要通过父框架找到按钮）
        for widget in self.history_combobox.master.master.winfo_children():
            if isinstance(widget, tk.Button) and widget.cget("text") == "保存录入内容":
                widget.config(state="disabled")
                break
        
        # 禁用删除记录按钮
        for widget in self.history_combobox.master.winfo_children():
            if isinstance(widget, tk.Button) and widget.cget("text") == "删除记录":
                widget.config(state="disabled")
                break
        
        # 禁用生成文档按钮
        for widget in self.history_combobox.master.master.winfo_children():
            if isinstance(widget, tk.Frame):  # 控制按钮区域
                for child in widget.winfo_children():
                    if isinstance(child, tk.Button) and child.cget("text") == "生成文档":
                        child.config(state="disabled")
                    elif isinstance(child, tk.Button) and child.cget("text") == "合并为PDF":
                        child.config(state="disabled")
                break
    
    def load_scheme_for_main(self, scheme_name):
        """
        为主操作界面加载方案
        """
        # 读取方案数据
        try:
            with open("app_data.json", "r", encoding="utf-8") as f:
                data = json.load(f)
            
            schemes = data.get("schemes", {})
            if scheme_name not in schemes:
                print(f"错误: 方案 '{scheme_name}' 不存在")
                return
            
            scheme_data = schemes[scheme_name]
            
            # 应用方案数据
            self.template_files = scheme_data.get("template_files", [])
            self.ordered_placeholders = scheme_data.get("placeholder_order", [])
            self.current_scheme = scheme_name
            
            # 创建输入字段
            self.create_input_fields()
            
            # 加载该方案的上次用户输入
            if self.current_scheme:
                last_inputs = self.load_user_inputs_for_scheme(self.current_scheme)
                if last_inputs:
                    # 填充上次的用户输入
                    for placeholder, value in last_inputs.items():
                        if placeholder in self.input_fields:
                            self.input_fields[placeholder].delete(0, tk.END)
                            self.input_fields[placeholder].insert(0, value)
            
        except Exception as e:
            print(f"错误: 加载方案时出错: {e}")
    
    def load_saved_schemes(self):
        """
        加载已保存的方案到列表框（保留此方法以保持向后兼容）
        """
        self.scheme_listbox_main.delete(0, tk.END)
        if os.path.exists("app_data.json"):
            try:
                with open("app_data.json", "r", encoding="utf-8") as f:
                    data = json.load(f)
                schemes = data.get("schemes", {})
                for scheme_name in schemes:
                    self.scheme_listbox_main.insert(tk.END, scheme_name)
            except Exception as e:
                print(f"加载方案时出错: {e}")
    
    def load_saved_schemes_combobox(self):
        """
        加载已保存方案下拉菜单
        """
        if os.path.exists("app_data.json"):
            try:
                with open("app_data.json", "r", encoding="utf-8") as f:
                    data = json.load(f)
                schemes = data.get("schemes", {})
                scheme_names = list(schemes.keys())
                self.saved_schemes_combobox['values'] = scheme_names
            except Exception as e:
                print(f"加载方案下拉菜单时出错: {e}")
    
    def setup_config_tab(self):
        """
        设置配置方案标签页（改为2栏式布局）
        """
        config_frame = ttk.Frame(self.config_frame, padding="10")
        config_frame.pack(fill=tk.BOTH, expand=True)
        
        # 配置网格权重
        config_frame.columnconfigure(0, weight=1)  # 左栏：方案配置

        config_frame.columnconfigure(1, weight=1)  # 右栏：用户录入
        config_frame.rowconfigure(1, weight=1)
        
        # 第一行：方案名称输入
        name_frame = ttk.Frame(config_frame)
        name_frame.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))
        name_frame.columnconfigure(1, weight=1)
        
        ttk.Label(name_frame, text="方案名称:").grid(row=0, column=0, sticky=tk.W)
        self.scheme_name_entry = ttk.Entry(name_frame)
        self.scheme_name_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=(5, 5))
        
        # 已保存方案下拉菜单
        self.saved_schemes_combobox = ttk.Combobox(name_frame, state="readonly", width=30)
        self.saved_schemes_combobox.grid(row=0, column=2, padx=(5, 5))
        self.saved_schemes_combobox.bind("<<ComboboxSelected>>", self.on_config_scheme_selected)
        
        # 删除选中方案按钮
        ttk.Button(name_frame, text="删除选中方案", command=self.delete_scheme).grid(row=0, column=3)
        
        # 两栏布局：模板文件 | 用户录入
        # 左栏：模板文件列表
        template_frame = ttk.LabelFrame(config_frame, text="模板文件", padding="10")
        template_frame.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(0, 5))
        
        # 右栏：用户录入区域
        input_frame = ttk.LabelFrame(config_frame, text="用户录入区预览", padding="10")
        input_frame.grid(row=1, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(5, 0))
        
        # 配置各栏的网格权重
        template_frame.columnconfigure(0, weight=1)
        template_frame.rowconfigure(0, weight=1)
        template_frame.rowconfigure(1, weight=0)  # 按钮区域不扩展
        
        input_frame.columnconfigure(0, weight=1)
        input_frame.rowconfigure(0, weight=1)
        
        # 左栏内容：模板文件列表
        self.config_template_listbox = tk.Listbox(template_frame, height=15, selectmode=tk.EXTENDED)
        self.config_template_listbox.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))
        
        # 移除了重复的scheme_listbox定义，因为它在配置方案界面中不需要
        
        # 创建一个框架来容纳按钮并使其填充整行
        button_frame = ttk.Frame(template_frame)
        button_frame.grid(row=1, column=0, columnspan=2, pady=5, sticky=(tk.W, tk.E))
        button_frame.columnconfigure(0, weight=1)
        button_frame.columnconfigure(1, weight=1)
        button_frame.columnconfigure(2, weight=1)
        
        # 创建3个按钮，使其填充整行并均匀分布
        ttk.Button(button_frame, text="添加模板文件", command=self.config_add_template_files).grid(row=0, column=0, sticky=(tk.W, tk.E), padx=2)
        ttk.Button(button_frame, text="移除选中文件", command=self.config_remove_selected_files).grid(row=0, column=1, sticky=(tk.W, tk.E), padx=2)
        ttk.Button(button_frame, text="检测占位符", command=self.config_detect_placeholders).grid(row=0, column=2, sticky=(tk.W, tk.E), padx=2)
        
        # 右栏内容：用户录入区域
        self.config_input_canvas = tk.Canvas(input_frame, height=200, relief=tk.SUNKEN, borderwidth=1)
        self.config_input_scrollbar = ttk.Scrollbar(input_frame, orient="vertical", command=self.config_input_canvas.yview)
        self.config_input_scrollable_frame = ttk.Frame(self.config_input_canvas)

        # 添加鼠标滚轮支持
        def _on_mousewheel(event):
            self.config_input_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        
        self.config_input_canvas.bind("<MouseWheel>", _on_mousewheel)
        self.config_input_scrollable_frame.bind("<MouseWheel>", _on_mousewheel)
        
        self.config_input_scrollable_frame.bind(
            "<Configure>",
            lambda e: self.config_input_canvas.configure(
                scrollregion=self.config_input_canvas.bbox("all")
            )
        )
        
        self.config_input_canvas.create_window((0, 0), window=self.config_input_scrollable_frame, anchor="nw")
        self.config_input_canvas.configure(yscrollcommand=self.config_input_scrollbar.set)
        
        self.config_input_canvas.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.config_input_scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        
        # 操作按钮区域
        button_frame = ttk.Frame(config_frame)
        button_frame.grid(row=2, column=0, columnspan=2, pady=(10, 0))
        
        ttk.Button(button_frame, text="保存当前方案", command=self.save_scheme).pack()
        
        # 初始化已保存方案下拉菜单
        self.load_saved_schemes_combobox()
    
    def update_config_input_area(self):
        """
        根据选中的方案更新用户录入区域
        """
        # 清空现有内容
        for widget in self.config_input_scrollable_frame.winfo_children():
            widget.destroy()
        
        # 重新生成输入框
        self.placeholder_entries = {}
        
        # 按顺序生成占位符输入框
        for i, placeholder in enumerate(self.ordered_placeholders):
            frame = ttk.Frame(self.config_input_scrollable_frame)
            frame.pack(fill=tk.X, pady=2)
            
            label = ttk.Label(frame, text=f"{placeholder}:")
            label.pack(side=tk.LEFT, padx=(0, 5))
            
            entry = ttk.Entry(frame)
            entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
            
            self.placeholder_entries[placeholder] = entry
    
    def save_scheme(self):
        """
        保存当前配置的方案
        """
        scheme_name = self.scheme_name_entry.get().strip()
        if not scheme_name:
            self.log_and_status("错误: 请输入方案名称")
            return
        
        if not self.template_files:
            self.log_and_status("错误: 请至少选择一个模板文件")
            return
        
        try:
            # 读取现有数据
            data = {}
            if os.path.exists("app_data.json"):
                with open("app_data.json", "r", encoding="utf-8") as f:
                    data = json.load(f)
            
            # 确保schemes键存在
            if "schemes" not in data:
                data["schemes"] = {}
            
            # 保存方案
            data["schemes"][scheme_name] = {
                "template_files": self.template_files.copy(),
                "placeholder_order": self.ordered_placeholders.copy()
            }
            
            # 写入文件
            with open("app_data.json", "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            
            # 更新下拉菜单
            self.load_saved_schemes_combobox()
            self.load_saved_schemes()
            
            self.log_and_status(f"成功: 方案 '{scheme_name}' 已保存")
        except Exception as e:
            self.log_and_status(f"错误: 保存方案时出错: {e}")
    
    def delete_scheme(self):
        """
        删除选中的方案
        """
        # 从下拉菜单中获取选中的方案名称
        scheme_name = self.saved_schemes_combobox.get()
        if not scheme_name:
            self.log_and_status("警告: 请先选择一个方案")
            return
        
        # 添加确认对话框
        from tkinter import messagebox
        result = messagebox.askyesno(
            "确认删除", 
            f"确定要删除方案 '{scheme_name}' 吗？此操作不可恢复。"
        )
        
        # 如果用户选择"否"，则取消删除操作
        if not result:
            return
        
        try:
            # 读取现有数据
            if not os.path.exists("app_data.json"):
                self.log_and_status("错误: 配置文件不存在")
                return
                
            with open("app_data.json", "r", encoding="utf-8") as f:
                data = json.load(f)
            
            schemes = data.get("schemes", {})
            if scheme_name not in schemes:
                self.log_and_status(f"错误: 方案 '{scheme_name}' 不存在")
                return
            
            # 删除方案
            del schemes[scheme_name]
            data["schemes"] = schemes
            
            # 保存更改
            with open("app_data.json", "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            
            # 更新列表框
            self.load_saved_schemes()
            self.load_saved_schemes_combobox()
            
            # 清空当前界面的内容
            self.scheme_name_entry.delete(0, tk.END)
            self.config_template_listbox.delete(0, tk.END)
            self.template_files.clear()
            
            # 清空用户录入区域
            for widget in self.config_input_scrollable_frame.winfo_children():
                widget.destroy()
                
            # 清空下拉菜单的选择
            self.saved_schemes_combobox.set('')
            
            self.log_and_status(f"成功: 方案 '{scheme_name}' 已删除")
        except Exception as e:
            self.log_and_status(f"错误: 删除方案时出错: {e}")
    
    def load_scheme(self):
        """
        加载选中的方案（配置方案界面）
        """
        selection = self.scheme_listbox.curselection()
        if not selection:
            print("警告: 请先选择一个方案")
            return

        selected_scheme = self.scheme_listbox.get(selection[0])
        try:
            # 读取方案数据
            with open("app_data.json", "r", encoding="utf-8") as f:
                data = json.load(f)
            
            schemes = data.get("schemes", {})
            if selected_scheme not in schemes:
                print(f"错误: 方案 '{selected_scheme}' 不存在")
                return
            
            scheme_data = schemes[selected_scheme]
            
            # 应用方案数据到配置界面
            self.template_files = scheme_data.get("template_files", [])
            self.ordered_placeholders = scheme_data.get("placeholder_order", [])
            
            # 更新模板文件列表
            self.config_template_listbox.delete(0, tk.END)
            for file_path in self.template_files:
                self.config_template_listbox.insert(tk.END, os.path.basename(file_path))
            
            # 更新占位符列表（如果存在）
            if hasattr(self, 'placeholder_listbox'):
                self.placeholder_listbox.delete(0, tk.END)
                for placeholder in self.ordered_placeholders:
                    if placeholder != '日期':  # 不显示日期占位符
                        self.placeholder_listbox.insert(tk.END, placeholder)
            
            # 更新方案名称
            self.scheme_name_entry.delete(0, tk.END)
            self.scheme_name_entry.insert(0, selected_scheme)
            
            # 创建用户输入控件
            self.config_create_input_fields()
            
            print(f"成功加载方案: {selected_scheme}")
        except Exception as e:
            print(f"加载方案时出错: {e}")
    
    def on_config_scheme_select(self, event):
        """
        配置方案界面中，当用户选择方案时的处理函数
        """
        selection = self.scheme_listbox.curselection()
        if not selection:
            return
        
        # 自动加载选中的方案
        self.load_scheme()
    
    def on_config_scheme_selected(self, event):
        """
        当选择方案时加载方案数据
        """
        selected = self.saved_schemes_combobox.get()
        if selected:  # 只有当有实际选择时才加载方案
            self.load_scheme_by_name(selected)
    
    def load_scheme_by_name(self, scheme_name):
        """
        根据方案名称加载方案
        :param scheme_name: 方案名称
        """
        try:
            if os.path.exists("app_data.json"):
                with open("app_data.json", "r", encoding="utf-8") as f:
                    data = json.load(f)
                
                schemes = data.get("schemes", {})
                if scheme_name in schemes:
                    scheme_data = schemes[scheme_name]
                    
                    # 清空当前内容
                    self.config_template_listbox.delete(0, tk.END)
                    self.template_files.clear()
                    
                    # 清空用户录入区域
                    for widget in self.config_input_scrollable_frame.winfo_children():
                        widget.destroy()
                    
                    # 加载模板文件
                    self.template_files.extend(scheme_data.get("template_files", []))
                    for file_path in self.template_files:
                        self.config_template_listbox.insert(tk.END, os.path.basename(file_path))
                    
                    # 加载占位符顺序
                    self.ordered_placeholders = scheme_data.get("placeholder_order", [])
                    
                    # 创建用户输入控件
                    self.config_create_input_fields()
                    
                    # 在方案名称输入框中显示当前方案名称
                    self.scheme_name_entry.delete(0, tk.END)
                    self.scheme_name_entry.insert(0, scheme_name)
                    
                    self.log_and_status(f"成功: 方案 '{scheme_name}' 已加载到配置界面")
                else:
                    self.log_and_status(f"错误: 方案 '{scheme_name}' 不存在")
            else:
                self.log_and_status("错误: 方案文件不存在")
        except Exception as e:
            self.log_and_status(f"错误: 加载方案时出错: {e}")
    
    def config_add_template_files(self):
        """
        添加模板文件（配置方案界面）
        """
        # 获取上次使用的模板目录
        last_template_dir = self.load_last_template_dir()
        
        files = filedialog.askopenfilenames(
            title="选择文档模板",
            filetypes=[
                ("所有文件", "*.*"),
                ("Word文档", "*.docx"),
                ("Excel表格", "*.xlsx")
            ],
            initialdir=last_template_dir
        )
        
        if files:
            # 保存所选文件的目录（使用第一个文件的目录）
            file_dir = os.path.dirname(files[0])
            self.save_last_template_dir(file_dir)
            
            for file in files:
                if file not in self.template_files:
                    self.template_files.append(file)
                    self.config_template_listbox.insert(tk.END, os.path.basename(file))

    def config_remove_selected_files(self):
        """
        移除选中的模板文件（配置方案界面）
        """
        # 获取所有选中的项目索引（从高到低排序，从后往前删除避免索引变化）
        selections = self.config_template_listbox.curselection()
        if not selections:
            print("请先选择要移除的文件")
            return
        
        # 从后往前删除，避免索引变化导致的问题
        for i in reversed(range(len(selections))):
            index = selections[i]
            # 从列表框中删除
            self.config_template_listbox.delete(index)
            # 从模板文件列表中删除
            if 0 <= index < len(self.template_files):
                del self.template_files[index]
    
    def config_detect_placeholders(self):
        """
        检测占位符（配置方案界面）
        """
        if not self.template_files:
            print("警告: 请先选择模板文件")
            return
        
        # 收集所有占位符和占位符文件映射
        result = self.processor.collect_all_placeholders(self.template_files)
        if isinstance(result, tuple) and len(result) == 2:
            self.placeholders, self.placeholder_files = result
        else:
            self.placeholders = result
            self.placeholder_files = {}
        
        # 创建有序的占位符列表（排除日期）
        # 如果已经有排序，则优先使用已有的排序，新增的占位符默认添加到最后面
        current_placeholders_set = set(self.ordered_placeholders)
        new_placeholders = self.placeholders - current_placeholders_set - {'日期'}
        
        # 将新增的占位符添加到现有排序列表的末尾
        self.ordered_placeholders.extend(sorted(new_placeholders))
        
        # 确保所有检测到的占位符都在列表中（防止某些占位符被意外删除）
        all_placeholders_except_date = [p for p in self.placeholders if p != '日期']
        existing_placeholders = set(self.ordered_placeholders)
        
        # 检查是否有缺失的占位符（理论上不应该有）
        missing_placeholders = set(all_placeholders_except_date) - existing_placeholders
        if missing_placeholders:
            self.ordered_placeholders.extend(sorted(missing_placeholders))
        
        # 移除已不存在的占位符
        self.ordered_placeholders = [p for p in self.ordered_placeholders if p in all_placeholders_except_date or p == '日期']
        
        # 创建输入字段
        self.config_create_input_fields()
    
    def move_up(self, index):
        """
        将指定索引的占位符向上移动
        :param index: 要移动的占位符索引
        """
        if index <= 0 or index >= len(self.ordered_placeholders):
            return
        
        # 交换位置
        self.ordered_placeholders[index], self.ordered_placeholders[index-1] = \
            self.ordered_placeholders[index-1], self.ordered_placeholders[index]
        
        # 重新创建输入界面
        self.config_create_input_fields()
    
    def config_create_input_fields(self):
        """
        创建输入字段（配置方案界面）
        """
        # 清除现有控件
        for widget in self.config_input_scrollable_frame.winfo_children():
            widget.destroy()
        
        # 创建输入字段和上移按钮
        self.input_fields = {}  # 用于存储输入框引用
        for i, placeholder in enumerate(self.ordered_placeholders):
            # 标签
            ttk.Label(self.config_input_scrollable_frame, text=f"{placeholder}:").grid(row=i, column=0, sticky=tk.W, pady=2)
            
            # 获取占位符配置
            config = self.get_placeholder_config(placeholder)
            
            # 根据配置创建不同类型的输入控件预览
            if config.get("type") == "combobox":
                # 创建下拉框预览
                options = config.get("options", [f"<{placeholder}>"])
                combobox = ttk.Combobox(self.config_input_scrollable_frame, values=options, width=25, state="readonly")
                combobox.grid(row=i, column=1, sticky=(tk.W, tk.E), pady=2, padx=(5, 0))
                combobox.set(options[0] if options else f"<{placeholder}>")
                self.input_fields[placeholder] = combobox
            else:
                # 创建普通文本框预览
                entry = ttk.Entry(self.config_input_scrollable_frame, width=25)
                entry.grid(row=i, column=1, sticky=(tk.W, tk.E), pady=2, padx=(5, 0))
                entry.insert(0, f"<{placeholder}>")
                entry.configure(state='readonly')  # 只读状态
                self.input_fields[placeholder] = entry
            
            # 设置按钮（在↑箭头左边）
            # 使用functools.partial来正确传递参数
            import functools
            setting_button = ttk.Button(
                self.config_input_scrollable_frame,
                text="⚙",
                width=3,
                command=functools.partial(self.configure_placeholder_type, placeholder)
            )
            setting_button.grid(row=i, column=2, pady=2, padx=(5, 0))
            
            # 上移按钮（第一个元素不显示）
            if i > 0:
                # 同样使用functools.partial来正确传递索引参数
                up_button = ttk.Button(
                    self.config_input_scrollable_frame, 
                    text="↑", 
                    width=3,
                    command=functools.partial(self.move_up, i)
                )
                up_button.grid(row=i, column=3, pady=2, padx=(5, 0))
        
        # 添加日期字段（自动生成，仅显示不提供输入）
        date_row = len(self.ordered_placeholders)
        ttk.Label(self.config_input_scrollable_frame, text="日期（自动生成）:").grid(row=date_row, column=0, sticky=tk.W, pady=2)
        date_label = ttk.Label(self.config_input_scrollable_frame, text=datetime.now().strftime('%Y年%m月%d日'))
        date_label.grid(row=date_row, column=1, sticky=tk.W, pady=2, padx=(5, 0))
        
        # 配置输入区域的列权重
        self.config_input_scrollable_frame.columnconfigure(1, weight=1)
        
        # 加载该方案的上次用户输入以正确显示下拉框选项
        if self.current_scheme:
            last_inputs = self.load_user_inputs_for_scheme(self.current_scheme)
            if last_inputs:
                # 更新下拉框预览中的选项和选中值
                for placeholder, value in last_inputs.items():
                    if placeholder in self.input_fields:
                        widget = self.input_fields[placeholder]
                        if isinstance(widget, ttk.Combobox):
                            # 对于下拉框，检查值是否在选项中，如果不在则添加
                            current_values = list(widget['values'])
                            if value not in current_values and value:
                                # 添加新值到选项中
                                current_values.append(value)
                                widget['values'] = current_values
                            # 设置当前值
                            widget.set(value)
    
    def configure_placeholder_type(self, placeholder):
        """
        配置占位符的输入框类型
        :param placeholder: 占位符名称
        """
        # 创建配置对话框
        config_dialog = tk.Toplevel(self.root)
        config_dialog.title(f"配置占位符 '{placeholder}'")
        config_dialog.geometry("300x250")
        config_dialog.resizable(False, False)
        
        # 居中显示对话框
        self.center_dialog(config_dialog, 300, 250)
        config_dialog.transient(self.root)
        config_dialog.grab_set()
        
        # 设置窗口图标
        self.set_dialog_icon(config_dialog)
        
        # 获取当前占位符配置（如果有的话）
        current_config = self.get_placeholder_config(placeholder)
        current_type = current_config.get("type", "entry")  # 默认为普通录入框
        current_options = current_config.get("options", [])
        
        # 类型选择
        ttk.Label(config_dialog, text="输入框类型:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        
        type_var = tk.StringVar(value=current_type)
        type_frame = ttk.Frame(config_dialog)
        type_frame.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=5, pady=5)
        
        entry_radio = ttk.Radiobutton(type_frame, text="普通录入框", variable=type_var, value="entry")
        entry_radio.pack(anchor=tk.W)
        
        combobox_radio = ttk.Radiobutton(type_frame, text="下拉框", variable=type_var, value="combobox")
        combobox_radio.pack(anchor=tk.W)
        
        date_radio = ttk.Radiobutton(type_frame, text="日期类型框", variable=type_var, value="date")
        date_radio.pack(anchor=tk.W)
        
        # 下拉框选项输入区域
        options_frame = ttk.Frame(config_dialog)
        options_frame.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E), padx=5, pady=5)
        
        ttk.Label(options_frame, text="下拉选项 (格式: 选项1,选项2,选项3...)(使用英文逗号):").grid(row=0, column=0, sticky=tk.W)
    
        options_text = tk.Text(options_frame, height=3, width=30)
        options_text.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=5)
        
        # 如果当前是下拉框类型，填充已有选项
        if current_type == "combobox" and current_options:
            options_text.insert("1.0", ",".join(current_options))
        
        # 根据选择的类型启用/禁用选项输入框
        def on_type_change(*args):
            if type_var.get() == "combobox":
                options_text.config(state="normal", bg="white")
            else:
                options_text.config(state="disabled", bg="#f0f0f0")
        
        type_var.trace("w", on_type_change)
        on_type_change()  # 初始化状态
        
        # 按钮区域
        button_frame = ttk.Frame(config_dialog)
        button_frame.grid(row=3, column=0, columnspan=2, pady=5)
        
        def save_config():
            # 保存配置
            config = {
                "type": type_var.get()
            }
            
            if config["type"] == "combobox":
                # 解析选项，支持逗号分隔
                options_text_content = options_text.get("1.0", tk.END).strip()
                if options_text_content:
                    # 按逗号分割并去除空格
                    options = [opt.strip() for opt in options_text_content.split(",") if opt.strip()]
                    config["options"] = options
                else:
                    config["options"] = []
            
            self.save_placeholder_config(placeholder, config)
            config_dialog.destroy()
            # 重新创建输入字段以反映更改
            self.config_create_input_fields()
        
        ttk.Button(button_frame, text="确定", command=save_config).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="取消", command=config_dialog.destroy).pack(side=tk.LEFT, padx=5)
    
    def create_input_fields(self):
        """
        创建输入字段（主操作界面）
        """
        # 清除现有控件
        for widget in self.input_scrollable_frame.winfo_children():
            widget.destroy()
        
        # 创建输入字段（移除上移按钮）
        self.input_fields = {}
        for i, placeholder in enumerate(self.ordered_placeholders):
            # 标签
            ttk.Label(self.input_scrollable_frame, text=f"{placeholder}:").grid(row=i, column=0, sticky=tk.W, pady=2)
            
            # 获取占位符配置
            config = self.get_placeholder_config(placeholder)
            
            # 根据配置创建不同类型的输入控件
            if config.get("type") == "combobox":
                # 创建下拉框，宽度增加到原来的1.5倍 (25 -> 37)
                options = config.get("options", [])
                combobox = ttk.Combobox(self.input_scrollable_frame, values=options, width=37, state="readonly")
                combobox.grid(row=i, column=1, sticky=(tk.W, tk.E), pady=2, padx=(5, 0))
                self.input_fields[placeholder] = combobox
            elif config.get("type") == "date":
                # 创建日期选择框
                date_frame = ttk.Frame(self.input_scrollable_frame)
                date_frame.grid(row=i, column=1, sticky=(tk.W, tk.E), pady=2, padx=(5, 0))
                
                # 日期显示标签
                today = datetime.now().strftime('%Y年%m月%d日')
                date_label = ttk.Label(date_frame, text=today)
                date_label.pack(side=tk.LEFT)
                
                # 修改日期按钮
                def create_modify_command(label_widget):
                    def modify_command():
                        self.modify_placeholder_date(label_widget)
                    return modify_command
                
                modify_button = ttk.Button(date_frame, text="选择日期", command=create_modify_command(date_label), width=8)
                modify_button.pack(side=tk.LEFT, padx=(5, 0))
                
                self.input_fields[placeholder] = date_label
            else:
                # 创建普通文本框，宽度增加到原来的1.5倍 (25 -> 37)
                entry = ttk.Entry(self.input_scrollable_frame, width=37)
                entry.grid(row=i, column=1, sticky=(tk.W, tk.E), pady=2, padx=(5, 0))
                self.input_fields[placeholder] = entry
            
            # 注意：已移除上移按钮，简化用户界面
        
        # 添加日期字段（自动生成，可选择修改）
        date_row = len(self.ordered_placeholders)
        ttk.Label(self.input_scrollable_frame, text="日期（自动生成）:").grid(row=date_row, column=0, sticky=tk.W, pady=2)
        
        # 创建日期显示和修改区域
        date_frame = ttk.Frame(self.input_scrollable_frame)
        date_frame.grid(row=date_row, column=1, sticky=tk.W, pady=2, padx=(5, 0))
        
        # 日期显示标签
        today = datetime.now().strftime('%Y年%m月%d日')
        self.date_label = ttk.Label(date_frame, text=today)
        self.date_label.pack(side=tk.LEFT)
        
        # 修改日期按钮
        ttk.Button(date_frame, text="选择日期", command=self.modify_date, width=8).pack(side=tk.LEFT, padx=(5, 0))
        
        # 将日期存储在input_fields中
        self.input_fields['日期'] = today
        
        # 配置输入区域的列权重
        self.input_scrollable_frame.columnconfigure(1, weight=1)
        
        # 加载该方案的上次用户输入以正确显示下拉框选项
        if self.current_scheme:
            last_inputs = self.load_user_inputs_for_scheme(self.current_scheme)
            if last_inputs:
                # 更新下拉框中的选项和选中值
                for placeholder, value in last_inputs.items():
                    if placeholder in self.input_fields:
                        widget = self.input_fields[placeholder]
                        if isinstance(widget, ttk.Combobox):
                            # 对于下拉框，检查值是否在选项中，如果不在则添加
                            current_values = list(widget['values'])
                            if value not in current_values and value:
                                # 添加新值到选项中
                                current_values.append(value)
                                widget['values'] = current_values
                            # 设置当前值
                            widget.set(value)
        
    def modify_placeholder_date(self, date_label):
        """
        修改占位符日期功能
        :param date_label: 日期标签控件
        """
        # 创建日期修改对话框
        date_dialog = tk.Toplevel(self.root)

        # 居中显示对话框
        self.center_dialog(date_dialog, 300, 100)
        date_dialog.transient(self.root)
        date_dialog.grab_set()

        date_dialog.title("选择日期")
        date_dialog.geometry("300x100")
        date_dialog.resizable(False, False)
        
        # 获取当前日期值
        current_date = date_label.cget("text")
        
        # 解析当前日期
        try:
            current_datetime = datetime.strptime(current_date, '%Y年%m月%d日')
            current_year = current_datetime.year
            current_month = current_datetime.month
            current_day = current_datetime.day
        except ValueError:
            current_datetime = datetime.now()
            current_year = current_datetime.year
            current_month = current_datetime.month
            current_day = current_datetime.day

        # 创建主框架
        main_frame = ttk.Frame(date_dialog, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 第一行：年月日选择
        date_selection_frame = ttk.Frame(main_frame)
        date_selection_frame.pack(fill=tk.X, pady=(0, 10))
        
        # 年份选择
        ttk.Label(date_selection_frame, text="年:").pack(side=tk.LEFT)
        year_var = tk.StringVar(value=str(current_year))
        years = [str(year) for year in range(datetime.now().year - 5, datetime.now().year + 6)]  # 前后5年
        year_combobox = ttk.Combobox(date_selection_frame, textvariable=year_var, values=years, width=8, state="readonly")
        year_combobox.pack(side=tk.LEFT, padx=(2, 5))
        
        # 月份选择
        ttk.Label(date_selection_frame, text="月:").pack(side=tk.LEFT)
        month_var = tk.StringVar(value=str(current_month))
        months = [str(i) for i in range(1, 13)]
        month_combobox = ttk.Combobox(date_selection_frame, textvariable=month_var, values=months, width=6, state="readonly")
        month_combobox.pack(side=tk.LEFT, padx=(2, 5))

        # 日期选择
        ttk.Label(date_selection_frame, text="日:").pack(side=tk.LEFT)
        day_var = tk.StringVar(value=str(current_day))
        days = [str(i) for i in range(1, 32)]
        day_combobox = ttk.Combobox(date_selection_frame, textvariable=day_var, values=days, width=6, state="readonly")
        day_combobox.pack(side=tk.LEFT, padx=(2, 0))

        # 第二行：按钮区域
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X)
        
        def use_today():
            today = datetime.now().strftime('%Y年%m月%d日')
            date_label.config(text=today)
            # 更新存储的值
            for placeholder, widget in self.input_fields.items():
                if widget == date_label:
                    # 保持控件引用不变，只更新值
                    self.input_fields[placeholder] = today
                    break
            date_dialog.destroy()
        
        def confirm_date():
            try:
                year = int(year_var.get())
                month = int(month_var.get())
                day = int(day_var.get())
                
                # 验证日期有效性
                datetime(year, month, day)
                
                # 格式化为指定格式
                new_date = f"{year}年{month:02d}月{day:02d}日"
                
                # 更新显示
                date_label.config(text=new_date)
                # 更新存储的值
                for placeholder, widget in self.input_fields.items():
                    if widget == date_label:
                        # 保持控件引用不变，只更新值
                        self.input_fields[placeholder] = new_date
                        break
                
                date_dialog.destroy()
            except ValueError:
                # 显示错误信息
                from tkinter import messagebox
                messagebox.showerror("日期错误", "请选择有效的日期！")
        
        def update_days(*args):
            """根据年份和月份更新日期选项"""
            try:
                year = int(year_var.get())
                month = int(month_var.get())
                
                # 计算该月的天数
                if month in [1, 3, 5, 7, 8, 10, 12]:
                    max_day = 31
                elif month in [4, 6, 9, 11]:
                    max_day = 30
                else:  # 2月
                    # 判断是否为闰年
                    if (year % 4 == 0 and year % 100 != 0) or (year % 400 == 0):
                        max_day = 29
                    else:
                        max_day = 28
                
                # 更新日期选项
                days = [str(i) for i in range(1, max_day + 1)]
                day_combobox['values'] = days
                
                # 如果当前选择的日期超出了该月的最大日期，则调整为最大日期
                current_day = int(day_var.get())
                if current_day > max_day:
                    day_var.set(str(max_day))
                    
            except (ValueError, tk.TclError):
                # 如果输入不完整或无效，不进行更新
                pass
        
        # 绑定年份和月份选择事件
        year_var.trace('w', update_days)
        month_var.trace('w', update_days)
        
        # 初始化日期选项
        update_days()

        # 按钮居中放置
        ttk.Button(button_frame, text="今天", command=use_today).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(button_frame, text="确定", command=confirm_date).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(button_frame, text="取消", command=date_dialog.destroy).pack(side=tk.LEFT, padx=(0, 5))
        
        # 使按钮框架内容居中
        button_frame.pack_configure(anchor=tk.CENTER)
    
    def modify_date(self):
        """
        修改日期功能 - 使用两行布局，年月日为一行，按钮为一行
        并实现动态日期选择，根据年月动态调整可选日期
        """
        # 创建日期修改对话框
        date_dialog = tk.Toplevel(self.root)

        # 居中显示对话框
        self.center_dialog(date_dialog, 300, 100)
        date_dialog.transient(self.root)
        date_dialog.grab_set()

        date_dialog.title("选择日期")
        date_dialog.geometry("300x100")
        date_dialog.resizable(False, False)
        
        # 获取当前日期值
        current_date = self.input_fields.get('日期', datetime.now().strftime('%Y年%m月%d日'))
        
        # 解析当前日期
        try:
            current_datetime = datetime.strptime(current_date, '%Y年%m月%d日')
            current_year = current_datetime.year
            current_month = current_datetime.month
            current_day = current_datetime.day
        except ValueError:
            current_datetime = datetime.now()
            current_year = current_datetime.year
            current_month = current_datetime.month
            current_day = current_datetime.day

        # 创建主框架
        main_frame = ttk.Frame(date_dialog, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 第一行：年月日选择
        date_selection_frame = ttk.Frame(main_frame)
        date_selection_frame.pack(fill=tk.X, pady=(0, 10))
        
        # 年份选择
        year_var = tk.StringVar(value=str(current_year))
        years = [str(year) for year in range(datetime.now().year - 10, datetime.now().year + 11)]  # 前后10年
        year_combobox = ttk.Combobox(date_selection_frame, textvariable=year_var, values=years, width=8, state="readonly")
        year_combobox.pack(side=tk.LEFT, padx=(2, 5))
        ttk.Label(date_selection_frame, text="年").pack(side=tk.LEFT)
        
        # 月份选择
        month_var = tk.StringVar(value=str(current_month))
        months = [str(i) for i in range(1, 13)]
        month_combobox = ttk.Combobox(date_selection_frame, textvariable=month_var, values=months, width=6, state="readonly")
        month_combobox.pack(side=tk.LEFT, padx=(2, 5))
        ttk.Label(date_selection_frame, text="月").pack(side=tk.LEFT)

        # 日期选择
        day_var = tk.StringVar(value=str(current_day))
        days = [str(i) for i in range(1, 32)]
        day_combobox = ttk.Combobox(date_selection_frame, textvariable=day_var, values=days, width=6, state="readonly")
        day_combobox.pack(side=tk.LEFT, padx=(2, 0))
        ttk.Label(date_selection_frame, text="日").pack(side=tk.LEFT)

        # 第二行：按钮区域
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X)
        
        def use_today():
            today = datetime.now().strftime('%Y年%m月%d日')
            self.date_label.config(text=today)
            self.input_fields['日期'] = today
            date_dialog.destroy()
        
        def confirm_date():
            try:
                year = int(year_var.get())
                month = int(month_var.get())
                day = int(day_var.get())
                
                # 验证日期有效性
                datetime(year, month, day)
                
                # 格式化为指定格式
                new_date = f"{year}年{month:02d}月{day:02d}日"
                
                # 更新显示和存储
                self.date_label.config(text=new_date)
                self.input_fields['日期'] = new_date
                
                date_dialog.destroy()
            except ValueError:
                # 显示错误信息
                from tkinter import messagebox
                messagebox.showerror("日期错误", "请选择有效的日期！")
        
        def update_days(*args):
            """根据年份和月份更新日期选项"""
            try:
                year = int(year_var.get())
                month = int(month_var.get())
                
                # 计算该月的天数
                if month in [1, 3, 5, 7, 8, 10, 12]:
                    max_day = 31
                elif month in [4, 6, 9, 11]:
                    max_day = 30
                else:  # 2月
                    # 判断是否为闰年
                    if (year % 4 == 0 and year % 100 != 0) or (year % 400 == 0):
                        max_day = 29
                    else:
                        max_day = 28
                
                # 更新日期选项
                days = [str(i) for i in range(1, max_day + 1)]
                day_combobox['values'] = days
                
                # 如果当前选择的日期超出了该月的最大日期，则调整为最大日期
                current_day = int(day_var.get())
                if current_day > max_day:
                    day_var.set(str(max_day))
                    
            except (ValueError, tk.TclError):
                # 如果输入不完整或无效，不进行更新
                pass
        
        # 绑定年份和月份选择事件
        year_var.trace('w', update_days)
        month_var.trace('w', update_days)
        
        # 初始化日期选项
        update_days()

        # 按钮居中放置
        ttk.Button(button_frame, text="今天", command=use_today).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(button_frame, text="确定", command=confirm_date).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(button_frame, text="取消", command=date_dialog.destroy).pack(side=tk.LEFT, padx=(0, 5))
        
        # 使按钮框架内容居中
        button_frame.pack_configure(anchor=tk.CENTER)
    
    def save_user_inputs(self):
        """
        保存用户录入的内容
        """
        if not self.current_scheme:
            self.log_and_status("警告: 请先选择一个方案")
            return
        
        # 收集用户输入
        user_inputs = {}
        for placeholder, widget in self.input_fields.items():
            if placeholder == '日期':
                user_inputs[placeholder] = widget  # 日期字段是字符串
                continue
                    
            if isinstance(widget, ttk.Entry):
                user_inputs[placeholder] = widget.get()
            elif isinstance(widget, ttk.Combobox):
                user_inputs[placeholder] = widget.get()
            elif isinstance(widget, ttk.Label):
                # 日期标签控件
                user_inputs[placeholder] = widget.cget("text")
            else:
                # 对于已更新为字符串值的日期类型占位符，直接使用值
                user_inputs[placeholder] = widget
        
        # 确保日期字段存在
        if '日期' not in user_inputs:
            today = datetime.now().strftime('%Y年%m月%d日')
            user_inputs['日期'] = today
        
        # 保存当前用户输入到历史记录
        self.save_to_history(user_inputs)
        
        # 保存当前用户输入
        if self.current_scheme:
            self.save_user_inputs_for_scheme(self.current_scheme, user_inputs)
            self.log_and_status(f"成功: 已保存'{self.current_scheme}'方案的用户录入内容")
    
    def save_to_history(self, user_inputs):
        """
        保存用户输入到历史记录
        :param user_inputs: 用户输入字典
        """
        try:
            # 读取现有历史记录
            history_data = {}
            
            if os.path.exists("app_data.json"):
                with open("app_data.json", "r", encoding="utf-8") as f:
                    data = json.load(f)
                    history_data = data.get("history", {})
            
            # 确保当前方案在历史记录中存在
            if self.current_scheme not in history_data:
                history_data[self.current_scheme] = []
            
            # 添加时间戳到记录中
            record = user_inputs.copy()
            record["__timestamp__"] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            
            # 检查是否与最近的一条记录完全相同（除了时间戳）
            should_save = True
            if history_data[self.current_scheme]:
                latest_record = history_data[self.current_scheme][0].copy()
                # 移除时间戳字段进行比较
                latest_record.pop("__timestamp__", None)
                current_record = record.copy()
                current_record.pop("__timestamp__", None)
                
                # 如果内容完全相同，则不保存
                if latest_record == current_record:
                    should_save = False
            
            if should_save:
                # 将新记录添加到开头
                history_data[self.current_scheme].insert(0, record)
                
                # 限制最多保存10条记录
                if len(history_data[self.current_scheme]) > 20:
                    history_data[self.current_scheme] = history_data[self.current_scheme][:20]
                
                # 读取所有数据并更新历史记录部分
                if os.path.exists("app_data.json"):
                    with open("app_data.json", "r", encoding="utf-8") as f:
                        all_data = json.load(f)
                else:
                    all_data = {}
                    
                all_data["history"] = history_data
                
                # 保存历史记录
                with open("app_data.json", "w", encoding="utf-8") as f:
                    json.dump(all_data, f, ensure_ascii=False, indent=2)
                
                # 更新下拉框内容
                self.update_history_combobox()
            
        except Exception as e:
            print(f"保存历史记录时出错: {e}")
    
    def update_history_combobox(self):
        """
        更新历史记录下拉框内容
        """
        try:
            if not os.path.exists("app_data.json"):
                self.history_combobox['values'] = []
                return
            
            with open("app_data.json", "r", encoding="utf-8") as f:
                data = json.load(f)
            
            history_data = data.get("history", {})
            
            # 获取当前方案的历史记录
            if self.current_scheme in history_data:
                # 创建显示文本列表
                history_texts = []
                for record in history_data[self.current_scheme]:
                    # 使用第一个非空的录入内容作为标识
                    display_info = "未命名记录"
                    # 遍历录入字段，找到第一个非空的内容
                    for key in record:
                        # 跳过时间戳和日期字段
                        if key not in ["__timestamp__", "日期"] and record[key]:
                            display_info = record[key]
                            break
                    history_texts.append(display_info)
                
                self.history_combobox['values'] = history_texts
            else:
                self.history_combobox['values'] = []
                
        except Exception as e:
            print(f"更新历史记录下拉框时出错: {e}")
    
    def load_history_record(self, event=None):
        """
        加载选中的历史记录到输入框
        :param event: 事件对象
        """
        try:
            # 获取选中项的索引
            selected_index = self.history_combobox.current()
            if selected_index < 0:
                return
            
            # 读取历史记录
            if not os.path.exists("app_data.json"):
                return
            
            with open("app_data.json", "r", encoding="utf-8") as f:
                data = json.load(f)
            
            history_data = data.get("history", {})
            
            # 获取当前方案的历史记录
            if self.current_scheme not in history_data:
                return
            
            # 获取选中的记录
            if selected_index >= len(history_data[self.current_scheme]):
                return
            
            record = history_data[self.current_scheme][selected_index]
            
            # 填充到输入框
            for placeholder, value in record.items():
                # 跳过时间戳字段
                if placeholder == "__timestamp__":
                    continue
                
                # 查找对应的输入控件
                if placeholder in self.input_fields:
                    widget = self.input_fields[placeholder]
                    if isinstance(widget, ttk.Entry):
                        widget.delete(0, tk.END)
                        widget.insert(0, value)
                    elif isinstance(widget, ttk.Combobox):
                        # 检查值是否在选项中，如果不在则添加
                        current_values = list(widget['values'])
                        if value not in current_values and value:
                            # 添加新值到选项中
                            current_values.append(value)
                            widget['values'] = current_values
                        # 设置当前值
                        widget.set(value)
                    elif isinstance(widget, ttk.Label):  # 日期标签
                        widget.config(text=value)
            
            self.log_and_status("已加载历史记录")
            
        except Exception as e:
            self.log_and_status(f"加载历史记录时出错: {e}")
    
    def generate_documents(self):
        """
        生成文档（在新线程中执行）
        """
        # 在新线程中执行生成文档操作
        thread = threading.Thread(target=self._generate_documents_thread)
        thread.daemon = True  # 设置为守护线程，确保主程序退出时线程也会退出
        thread.start()
    
    def _generate_documents_thread(self):
        """
        在线程中执行文档生成操作
        """
        
        # 处理模板
        try:
            self.update_status("开始生成文档...")
            # 收集用户输入
            user_inputs = {}
            for placeholder, widget in self.input_fields.items():
                if placeholder == '日期':
                    user_inputs[placeholder] = widget  # 日期字段是字符串
                    continue
                        
                if isinstance(widget, ttk.Entry):
                    user_inputs[placeholder] = widget.get()
                elif isinstance(widget, ttk.Combobox):
                    user_inputs[placeholder] = widget.get()
                elif isinstance(widget, ttk.Label):
                    # 日期标签控件
                    user_inputs[placeholder] = widget.cget("text")
            
            # 确保日期字段存在
            if '日期' not in user_inputs:
                today = datetime.now().strftime('%Y年%m月%d日')
                user_inputs['日期'] = today
            self.generated_files = self.processor.process_templates(self.template_files, user_inputs, self.output_dir)
            self.log_and_status(f"成功: 文档生成完成！文件已保存到 {self.output_dir} 目录中。")
            
            # 移除自动询问打开输出文件夹的功能
            # self.ask_to_open_output_dir()
        except Exception as e:
            self.log_and_status(f"错误: 生成文档时出错：{str(e)}")

    def merge_to_pdf(self):
        """
        将生成的文档合并为PDF（在新线程中执行）
        """
        # 在新线程中执行合并PDF操作
        thread = threading.Thread(target=self._merge_to_pdf_thread)
        thread.daemon = True  # 设置为守护线程，确保主程序退出时线程也会退出
        thread.start()
    
    def _merge_to_pdf_thread(self):
        """
        在线程中执行PDF合并操作
        """
        if not hasattr(self, 'generated_files') or not self.generated_files:
            self.log_and_status("警告: 请先生成文档")
            return
        
        # 检查是否有需要转换为PDF的文件
        docx_files = [f for f in self.generated_files if f.endswith('.docx')]
        xlsx_files = [f for f in self.generated_files if f.endswith('.xlsx')]
        
        if not docx_files and not xlsx_files:
            self.log_and_status("警告: 没有找到可转换为PDF的文档")
            return
        
        try:
            # 初始化文件状态跟踪
            all_files = [os.path.basename(f) for f in docx_files + xlsx_files]
            
            # 在主线程中创建进度窗口
            self.root.after(0, lambda: self._create_pdf_progress_window(len(all_files)))
            
            # 等待进度窗口创建完成
            import time
            time.sleep(0.1)
            
            # 初始化所有文件状态为等待
            for filename in all_files:
                if hasattr(self, 'pdf_progress_window'):
                    self.root.after(0, lambda f=filename: self.update_pdf_progress(f, "waiting"))
            
            pdf_files = []  # 初始化pdf_files列表
            
            # 将生成的Word文档转换为PDF
            if docx_files:
                self.update_status("开始转换Word文档为PDF...")
                # 设置进度回调
                self.processor.set_progress_callback(self._pdf_progress_callback)
                pdf_files = self.processor.convert_docx_to_pdf(docx_files, status_callback=self.update_status)
                self.update_status(f"成功转换 {len(pdf_files)} 个PDF文件")
            
            # 将生成的Excel文件转换为PDF
            if xlsx_files:
                self.update_status("开始转换Excel文件为PDF...")
                # 设置进度回调
                self.processor.set_progress_callback(self._pdf_progress_callback)
                xlsx_pdf_files = self.processor.convert_xlsx_to_pdf(xlsx_files, status_callback=self.update_status)
                pdf_files.extend(xlsx_pdf_files)
                self.update_status(f"成功转换Excel为PDF {len(xlsx_pdf_files)} 个文件")
            
            # 检查是否有成功转换的PDF文件
            if not pdf_files:
                self.log_and_status("警告: 没有成功转换为PDF的文件")
                return
            
            # 合并PDF文件
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            merged_pdf_path = os.path.join(self.output_dir, f"合并文档_{timestamp}.pdf")
            self.update_status(f"开始合并PDF文件到: {merged_pdf_path}")
            self.processor.merge_pdfs(pdf_files, merged_pdf_path, status_callback=self.update_status)
            
            # 清理单个PDF文件
            self.processor.cleanup_single_pdfs(pdf_files, status_callback=self.update_status)
            
            self.log_and_status(f"成功: PDF合并完成！文件已保存为: {os.path.basename(merged_pdf_path)}")
            
            # 启用关闭按钮
            if hasattr(self, 'pdf_progress_window'):
                self.root.after(0, lambda: self.close_progress_button.config(state=tk.NORMAL))
        except Exception as e:
            error_msg = f"错误: 合并PDF时出错：{str(e)}"
            self.log_and_status(error_msg)
            # 使用状态栏显示替代弹窗提示，符合用户偏好
            self.update_status("请重启软件并关闭所有office相关软件。")
    
    def _pdf_progress_callback(self, filename, status):
        """
        PDF转换进度回调函数
        :param filename: 文件名
        :param status: 状态
        """
        if hasattr(self, 'pdf_progress_window') and self.pdf_progress_window:
            self.root.after(0, lambda: self.update_pdf_progress(filename, status))
    
    def _create_pdf_progress_window(self, total_files):
        """
        创建PDF转换进度窗口
        :param total_files: 总文件数
        """
        # 创建进度窗口
        self.pdf_progress_window = tk.Toplevel(self.root)
        self.pdf_progress_window.title("PDF转换进度")
        self.pdf_progress_window.geometry("500x400")
        self.pdf_progress_window.resizable(False, False)
        
        # 居中显示
        self.center_dialog(self.pdf_progress_window, 500, 400)
        self.pdf_progress_window.transient(self.root)
        self.pdf_progress_window.grab_set()
        
        # 设置窗口图标
        self.set_dialog_icon(self.pdf_progress_window)
        
        # 创建主框架
        main_frame = ttk.Frame(self.pdf_progress_window, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 标题
        ttk.Label(main_frame, text="PDF转换进度", font=("Arial", 12, "bold")).pack(pady=(0, 10))
        
        # 进度信息
        self.progress_info_label = ttk.Label(main_frame, text=f"总共需要转换 {total_files} 个文件")
        self.progress_info_label.pack(anchor=tk.W)
        
        # 进度条
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(main_frame, variable=self.progress_var, maximum=total_files)
        self.progress_bar.pack(fill=tk.X, pady=(5, 10))
        
        # 当前任务标签
        self.current_task_label = ttk.Label(main_frame, text="准备开始转换...")
        self.current_task_label.pack(anchor=tk.W, pady=(0, 10))
        
        # 文件列表框架
        list_frame = ttk.LabelFrame(main_frame, text="文件转换状态", padding="5")
        list_frame.pack(fill=tk.BOTH, expand=True)
        
        # 创建列表框和滚动条
        listbox_frame = ttk.Frame(list_frame)
        listbox_frame.pack(fill=tk.BOTH, expand=True)
        
        self.file_listbox = tk.Listbox(listbox_frame, height=15)
        self.file_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        scrollbar = ttk.Scrollbar(listbox_frame, orient=tk.VERTICAL, command=self.file_listbox.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.file_listbox.config(yscrollcommand=scrollbar.set)
        
        # 关闭按钮（初始禁用）
        self.close_progress_button = ttk.Button(main_frame, text="关闭", command=self.pdf_progress_window.destroy, state=tk.DISABLED)
        self.close_progress_button.pack(pady=(10, 0))
        
        # 初始化文件状态跟踪
        self.file_status = {}
        self.completed_files = 0
        self.total_files = total_files
    
    def update_pdf_progress(self, filename, status):
        """
        更新PDF转换进度
        :param filename: 文件名
        :param status: 状态 (waiting, converting, completed, failed)
        """
        if not hasattr(self, 'pdf_progress_window') or not self.pdf_progress_window:
            return
        
        # 更新文件状态
        self.file_status[filename] = status
        
        # 更新列表框显示
        self._update_progress_display(filename, status)
    
    def _update_progress_display(self, filename, status):
        """
        更新进度显示
        :param filename: 文件名
        :param status: 状态
        """
        # 清空列表框
        self.file_listbox.delete(0, tk.END)
        
        # 重新添加所有文件及其状态
        status_symbols = {
            "waiting": "⏳",
            "converting": "🔄",
            "completed": "✅",
            "failed": "❌"
        }
        
        for file, stat in self.file_status.items():
            symbol = status_symbols.get(stat, "❓")
            display_text = f"{symbol} {file}"
            self.file_listbox.insert(tk.END, display_text)
            
            # 根据状态设置颜色
            if stat == "completed":
                self.file_listbox.itemconfig(tk.END, {'fg': 'green'})
            elif stat == "failed":
                self.file_listbox.itemconfig(tk.END, {'fg': 'red'})
            elif stat == "converting":
                self.file_listbox.itemconfig(tk.END, {'fg': 'blue'})
        
        # 更新进度条和信息
        if status == "completed":
            self.completed_files += 1
            
        self.progress_var.set(self.completed_files)
        self.progress_info_label.config(text=f"已完成: {self.completed_files}/{self.total_files}")
        
        # 更新当前任务标签
        status_texts = {
            "waiting": "等待转换",
            "converting": "正在转换",
            "completed": "转换完成",
            "failed": "转换失败"
        }
        self.current_task_label.config(text=f"正在处理: {filename} ({status_texts.get(status, '未知状态')})")
        
        # 自动滚动到当前处理的文件位置
        if status in ["converting", "completed", "failed"]:
            # 找到当前文件在列表中的位置
            for i, (file, stat) in enumerate(self.file_status.items()):
                if file == filename:
                    # 滚动到该位置
                    self.file_listbox.see(i)
                    # 选中该项以便更清楚地看到
                    self.file_listbox.selection_clear(0, tk.END)
                    self.file_listbox.selection_set(i)
                    break
        
        # 如果所有文件都处理完成，启用关闭按钮
        if self.completed_files >= self.total_files:
            self.close_progress_button.config(state=tk.NORMAL)
            self.current_task_label.config(text="所有文件转换完成！")

def main():
    """
    主函数
    """
    root = tk.Tk()
    app = DocumentProcessorUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()
