import os
import re
from datetime import datetime
from docx import Document
from openpyxl import Workbook, load_workbook

# 设置标志位为True，因为我们现在直接导入了这些模块
EXCEL_PROCESSING_AVAILABLE = True
PDF_CONVERSION_AVAILABLE = True
PDF_MERGING_AVAILABLE = True

class DocumentProcessor:
    def __init__(self):
        """
        初始化文档处理器
        """
        self.placeholders = set()  # 存储所有占位符
        self.user_inputs = {}  # 存储用户输入
        self.template_files = []  # 存储选中的模板文件
        self.progress_callback = None  # 进度回调函数

    def set_progress_callback(self, callback):
        """
        设置进度回调函数
        :param callback: 回调函数
        """
        self.progress_callback = callback

    def extract_placeholders_from_docx(self, file_path):
        """
        从Word文档中提取占位符
        :param file_path: Word文档路径
        :return: 占位符集合
        """
        placeholders = set()
        doc = Document(file_path)
        
        # 提取段落中的占位符
        for paragraph in doc.paragraphs:
            placeholders.update(self.find_placeholders_in_text(paragraph.text))
        
        # 提取表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    placeholders.update(self.find_placeholders_in_text(cell.text))
        
        return placeholders

    def extract_placeholders_from_xlsx(self, file_path):
        """
        从Excel文件中提取占位符
        :param file_path: Excel文件路径
        :return: 占位符集合
        """
        if not EXCEL_PROCESSING_AVAILABLE:
            raise Exception("Excel处理功能不可用，请安装openpyxl库")
        
        placeholders = set()
        workbook = load_workbook(file_path)
        
        # 遍历所有工作表
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            
            # 遍历所有单元格
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        placeholders.update(self.find_placeholders_in_text(cell.value))
        
        return placeholders

    def find_placeholders_in_text(self, text):
        """
        在文本中查找占位符
        :param text: 要搜索的文本
        :return: 占位符集合
        """
        # 使用正则表达式查找形如 {占位符名称} 的占位符
        pattern = r'\{([^}]+)\}'
        matches = re.findall(pattern, text)
        return set(matches)

    def collect_all_placeholders(self, template_files):
        """
        从所有模板文件中收集占位符
        :param template_files: 模板文件列表
        :return: 所有占位符的集合和占位符到文件的映射
        """
        all_placeholders = set()
        placeholder_files = {}  # 记录每个占位符出现在哪些文件中
        for file_path in template_files:
            if file_path.endswith('.docx'):
                placeholders = self.extract_placeholders_from_docx(file_path)
            elif file_path.endswith('.xlsx'):
                placeholders = self.extract_placeholders_from_xlsx(file_path)
            else:
                continue
                
            all_placeholders.update(placeholders)
            
            # 记录每个占位符出现的文件
            for placeholder in placeholders:
                if placeholder not in placeholder_files:
                    placeholder_files[placeholder] = []
                placeholder_files[placeholder].append(file_path)
                
        return all_placeholders, placeholder_files

    def replace_placeholders_in_docx(self, template_path, output_path, replacements):
        """
        在Word文档中替换占位符
        :param template_path: 模板文件路径
        :param output_path: 输出文件路径
        :param replacements: 替换字典
        """
        doc = Document(template_path)
        
        # 替换段落中的占位符
        for paragraph in doc.paragraphs:
            self.replace_text_in_paragraph(paragraph, replacements)
        
        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.replace_text_in_paragraph(paragraph, replacements)
        
        # 替换页眉页脚中的占位符
        for section in doc.sections:
            # 替换页眉中的占位符
            for paragraph in section.header.paragraphs:
                self.replace_text_in_paragraph(paragraph, replacements)
            
            # 替换页脚中的占位符
            for paragraph in section.footer.paragraphs:
                self.replace_text_in_paragraph(paragraph, replacements)
        
        # 保存新文档
        doc.save(output_path)

    def replace_placeholders_in_xlsx(self, template_path, output_path, replacements):
        """
        在Excel文件中替换占位符
        :param template_path: 模板文件路径
        :param output_path: 输出文件路径
        :param replacements: 替换字典
        """
        if not EXCEL_PROCESSING_AVAILABLE:
            raise Exception("Excel处理功能不可用，请安装openpyxl库")
        
        workbook = load_workbook(template_path)
        
        # 遍历所有工作表
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            
            # 遍历所有单元格
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value:
                        # 处理字符串类型的单元格
                        if isinstance(cell.value, str):
                            new_value = cell.value
                            for key, value in replacements.items():
                                placeholder = '{' + key + '}'
                                if placeholder in new_value:
                                    new_value = new_value.replace(placeholder, str(value))
                            cell.value = new_value
                        # 处理公式中的占位符（如果有的话）
                        elif isinstance(cell.value, (int, float)) and cell.data_type == 'f' and cell.value is not None:
                            # 处理公式类型
                            pass  # 公式中的占位符处理需要特殊处理，暂时跳过
        
        # 保存新文件
        workbook.save(output_path)

    def replace_text_in_paragraph(self, paragraph, replacements):
        """
        在段落中替换文本（改进版，更好地保持格式）
        :param paragraph: 段落对象
        :param replacements: 替换字典
        """
        # 遍历所有需要替换的占位符
        for placeholder, replacement in replacements.items():
            # 使用完整的占位符格式进行替换 {占位符}
            full_placeholder = f'{{{placeholder}}}'
            
            # 遍历段落中的每个run
            for run in paragraph.runs:
                if full_placeholder in run.text:
                    # 执行替换操作
                    run.text = run.text.replace(full_placeholder, str(replacement))

    def process_templates(self, template_files, user_inputs, output_dir):
        """
        处理模板文件，生成新文件
        :param template_files: 模板文件列表
        :param user_inputs: 用户输入字典
        :param output_dir: 输出目录
        :return: 生成的文件列表
        """
        generated_files = []
        total_files = len(template_files)
        
        # 确保输出目录存在
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # 处理每个模板文件
        for i, template_file in enumerate(template_files):
            try:
                # 调用进度回调函数
                if self.progress_callback:
                    self.progress_callback(i + 1, total_files, os.path.basename(template_file), "processing")
                
                # 生成输出文件路径
                filename = os.path.basename(template_file)
                name, ext = os.path.splitext(filename)
                # 在文件名后添加日期
                today = datetime.now().strftime('%Y%m%d')
                output_filename = f"{name}_{today}{ext}"
                output_path = os.path.join(output_dir, output_filename)
                
                # 检查目标文件是否已存在，如果存在则添加序号
                counter = 1
                while os.path.exists(output_path):
                    output_filename = f"{name}_{today}_{counter}{ext}"
                    output_path = os.path.join(output_dir, output_filename)
                    counter += 1
                
                # 根据文件类型处理文件
                if template_file.endswith('.docx'):
                    self.replace_placeholders_in_docx(template_file, output_path, user_inputs)
                elif template_file.endswith('.xlsx'):
                    self.replace_placeholders_in_xlsx(template_file, output_path, user_inputs)
                else:
                    # 对于不支持的文件类型，复制原文件
                    import shutil
                    shutil.copy2(template_file, output_path)
                
                generated_files.append(output_path)
                
                # 调用进度回调函数
                if self.progress_callback:
                    self.progress_callback(i + 1, total_files, os.path.basename(template_file), "completed")
                    
            except Exception as e:
                # 调用进度回调函数
                if self.progress_callback:
                    self.progress_callback(i + 1, total_files, os.path.basename(template_file), "failed")
                raise Exception(f"处理文件 {template_file} 时出错: {str(e)}")
        
        return generated_files